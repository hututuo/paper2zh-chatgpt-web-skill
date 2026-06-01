#!/usr/bin/env python3
"""Local JSON pipeline for manual ChatGPT Web document translation.

The script is intentionally mechanical. It never generates translation text.
It prepares JSON jobs, accepts downloaded JSON outputs, validates identifiers
and hashes, merges model-authored text, runs QA patch application, and renders
Markdown/DOCX outputs.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Pt, RGBColor
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception as exc:  # pragma: no cover - reported by commands.
    Document = None  # type: ignore[assignment]
    DocxDocument = object  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    CT_Tbl = object  # type: ignore[assignment]
    CT_P = object  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    RGBColor = None  # type: ignore[assignment]
    Table = object  # type: ignore[assignment]
    Paragraph = object  # type: ignore[assignment]
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
SCHEMA_VERSION = "1.0"
TRANSLATION_STAGE = "translation"
QA_STAGE = "qa"
DEFAULT_BODY_TRANSLATION_MAX_ITEMS = 100
DEFAULT_BODY_TRANSLATION_MAX_CHARS = 60000
DEFAULT_AUX_TRANSLATION_MAX_ITEMS = 500
DEFAULT_AUX_TRANSLATION_MAX_CHARS = 60000
DEFAULT_QA_MAX_ITEMS = 200000
DEFAULT_QA_MAX_CHARS = 2000000
RICH_TOKEN_PATTERN = re.compile(r"\[\[RT:[A-Za-z0-9_:\-]+\]\]")
INLINE_FORMAT_TAG_PATTERN = re.compile(r"</?(?:b|i|u|s|sup|sub)>")
RICH_TOKEN_OR_FORMAT_TAG_PATTERN = re.compile(
    r"(\[\[RT:[A-Za-z0-9_:\-]+\]\]|</?(?:b|i|u|s|sup|sub)>)"
)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_VAL = f"{{{W_NS}}}val"
CJK_HEADING_EXCLUDED_STARTS = (
    "本研究",
    "本文",
    "结果显示",
    "此外",
    "同时",
    "因此",
    "然而",
    "其中",
    "通过",
    "根据",
    "对于",
    "虽然",
    "在本",
    "为了",
    "随着",
    "由于",
    "如果",
    "当",
    "与",
    "其",
    "该",
    "综上",
)

AUXILIARY_SECTION_PATTERNS = (
    "references",
    "bibliography",
    "reference list",
    "supplement",
    "supplementary",
    "appendix",
    "acknowledg",
    "funding",
    "author contribution",
    "contributions",
    "conflict of interest",
    "competing interest",
    "data availability",
    "ethics",
    "附录",
    "补充",
    "参考文献",
    "致谢",
    "基金",
    "作者贡献",
    "利益冲突",
    "数据可用",
    "伦理",
)

LATEX_MAP = {
    r'\pm': '±', r'\times': '×', r'\div': '÷', r'\le': '≤', r'\ge': '≥',
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε',
    r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\omicron': 'ο',
    r'\pi': 'π', r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ',
    r'\phi': 'φ', r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω', r'\Delta': 'Δ',
    r'\infty': '∞', r'\approx': '≈', r'\neq': '≠', r'\cdot': '·',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\uparrow': '↑', r'\downarrow': '↓',
    r'\partial': '∂', r'\nabla': '∇', r'\sum': '∑', r'\prod': '∏', r'\coprod': '∐',
    r'\sim': '~', r'\equiv': '≡', r'\propto': '∝', r'\circ': '∘', r'\degree': '°',
    r'\circ C': '°C', r'^\circ': '°', r'^\circ C': '°C', r'\sqrt': '√', r'\int': '∫',
    r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\supset': '⊃', r'\subseteq': '⊆',
    r'\supseteq': '⊇', r'\cup': '∪', r'\cap': '∩', r'\emptyset': '∅',
    r'\forall': '∀', r'\exists': '∃', r'\angle': '∠', r'\perp': '⊥', r'\parallel': '∥',
    r'\simeq': '≃', r'\cong': '≅', r'\ast': '∗', r'\star': '★', r'\bullet': '•',
    r'\lneq': '≨', r'\gneq': '≩', r'\ll': '≪', r'\gg': '≫',
}

def _replace_math_symbols(paragraph: Any) -> int:
    if "$" not in paragraph.text:
        return 0
    replacement_count = 0
    for run in paragraph.runs:
        if "$" in run.text:
            text = run.text.replace("$", "")
            for k, v in sorted(LATEX_MAP.items(), key=lambda pair: len(pair[0]), reverse=True):
                if k in text:
                    replacement_count += text.count(k)
                    text = text.replace(k, v)
            while "  " in text:
                text = text.replace("  ", " ")
            run.text = text.strip()
    return replacement_count


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def fail(message: str) -> None:
    raise SystemExit(f"ERROR: {message}")


def ensure_docx_available() -> None:
    if DOCX_IMPORT_ERROR is not None:
        fail(f"python-docx is required but could not be imported: {DOCX_IMPORT_ERROR}")


def sha256_bytes(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def count_words(text: str) -> int:
    words = re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?|[\u4e00-\u9fff]", text)
    return len(words)


def load_json(path: Path) -> Any:
    text = path.read_text(encoding="utf-8")
    try:
        return json.loads(extract_json_text(text))
    except (json.JSONDecodeError, ValueError) as exc:
        raise ValueError(f"{path}: could not extract valid JSON object/array from model output") from exc


def extract_json_text(text: str) -> str:
    """Extract the first valid JSON object/array from model output text."""

    stripped = text.strip()
    fence = re.search(r"```(?:json[0-9a-zA-Z_-]*)?\s*(.*?)\s*```", stripped, re.S | re.I)
    if fence:
        candidate = fence.group(1).strip()
        try:
            json.loads(candidate)
            return candidate
        except json.JSONDecodeError:
            pass

    decoder = json.JSONDecoder()
    for idx, char in enumerate(stripped):
        if char not in "{[":
            continue
        try:
            _data, end = decoder.raw_decode(stripped[idx:])
        except json.JSONDecodeError:
            continue
        return stripped[idx : idx + end]
    try:
        json.loads(stripped)
    except json.JSONDecodeError as exc:
        raise ValueError("could not extract valid JSON object/array from model output") from exc
    return stripped


def load_json_obj(path: Path) -> dict[str, Any]:
    data = load_json(path)
    if not isinstance(data, dict):
        raise ValueError(f"{path} must contain a JSON object")
    return data


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def rel(path: Path, base: Path) -> str:
    try:
        return str(path.relative_to(base))
    except ValueError:
        return str(path)


def open_path(path: Path) -> None:
    if sys.platform == "darwin":
        subprocess.run(["open", str(path)], check=False)


def ensure_bundle_dirs(bundle: Path) -> None:
    for sub in [
        "jobs/translation/downloads",
        "jobs/qa/downloads",
        "logs",
        "assembled",
        "outputs",
    ]:
        (bundle / sub).mkdir(parents=True, exist_ok=True)


def read_prompt_template(name: str) -> str:
    return (SKILL_DIR / "prompts" / name).read_text(encoding="utf-8").strip()


def iter_docx_blocks(document: DocxDocument):
    body = document.element.body
    paragraph_index = 0
    table_index = 0
    docx_block_index = 0
    for child in body.iterchildren():
        docx_block_index += 1
        if isinstance(child, CT_P):
            paragraph_index += 1
            yield {
                "kind": "paragraph",
                "docx_block_index": docx_block_index,
                "paragraph_index": paragraph_index,
                "object": Paragraph(child, document),
            }
        elif isinstance(child, CT_Tbl):
            table_index += 1
            yield {
                "kind": "table",
                "docx_block_index": docx_block_index,
                "table_index": table_index,
                "object": Table(child, document),
            }


def classify_paragraph_info(text: str, style_name: str) -> tuple[str, str]:
    style = style_name.lower()
    if style.startswith("heading") or style in {"title", "subtitle"}:
        return "heading", "style"
    if "caption" in style:
        return "caption", "none"
    if len(text) <= 140 and text.isupper() and re.search(r"[A-Z]", text):
        return "heading", "heuristic"
    cleaned = re.sub(r"\s+", " ", text).strip()
    has_cjk = bool(re.search(r"[\u4e00-\u9fff]", cleaned))
    terminal = bool(re.search(r"[。！？.!?；;：:]$", cleaned))
    punctuation_count = len(re.findall(r"[，,。.!？?；;：:、]", cleaned))
    if (
        has_cjk
        and 2 <= len(cleaned) <= 25
        and not terminal
        and punctuation_count <= 1
        and not cleaned.startswith(CJK_HEADING_EXCLUDED_STARTS)
    ):
        return "heading", "heuristic"
    return "paragraph", "none"


def classify_paragraph(text: str, style_name: str) -> str:
    return classify_paragraph_info(text, style_name)[0]


def element_text_with_math(element: Any) -> str:
    """Collect visible text plus Office Math text nodes from an OOXML element."""

    parts: list[str] = []
    for node in element.iter():
        tag = str(getattr(node, "tag", ""))
        if tag.endswith("}t") or tag.endswith("}instrText"):
            if node.text:
                parts.append(node.text)
        elif tag.endswith("}tab"):
            parts.append("\t")
        elif tag.endswith("}br"):
            parts.append("\n")
    return "".join(parts)


def paragraph_text(paragraph: Paragraph) -> str:
    text = element_text_with_math(paragraph._p).strip()
    return text if text else paragraph.text.strip()


def cell_text(cell: Any) -> str:
    text = element_text_with_math(cell._tc).strip()
    return text if text else cell.text.strip()


def cell_key(cell: Any) -> str:
    # python-docx exposes merged-cell aliases as the same canonical tc path.
    # The XML path is stable across wrapper lifetimes, unlike Python object ids.
    return str(cell._tc.getroottree().getpath(cell._tc))


def local_name(element: Any) -> str:
    tag = str(getattr(element, "tag", ""))
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def is_math_element(element: Any) -> bool:
    tag = str(getattr(element, "tag", ""))
    return "officeDocument/2006/math" in tag


def has_rich_run_properties(run_element: Any) -> bool:
    rpr = run_element.find(f".//{{{W_NS}}}rPr")
    if rpr is None:
        return False
    rich_names = {
        "b",
        "i",
        "u",
        "strike",
        "dstrike",
        "vertAlign",
        "smallCaps",
        "caps",
        "position",
        "highlight",
        "color",
    }
    return any(local_name(child) in rich_names for child in rpr.iterchildren())


def word_bool_enabled(element: Any) -> bool:
    value = str(element.get(W_VAL, "")).lower()
    return value not in {"0", "false", "off", "none"}


def run_format_tags(run_element: Any) -> tuple[str, str]:
    rpr = run_element.find(f"{{{W_NS}}}rPr")
    if rpr is None:
        return "", ""
    open_parts: list[str] = []
    close_parts: list[str] = []
    seen: set[str] = set()
    for child in rpr.iterchildren():
        name = local_name(child)
        tag = ""
        if name == "vertAlign":
            value = str(child.get(W_VAL, ""))
            if value == "superscript":
                tag = "sup"
            elif value == "subscript":
                tag = "sub"
        elif name == "b" and word_bool_enabled(child):
            tag = "b"
        elif name == "i" and word_bool_enabled(child):
            tag = "i"
        elif name == "u" and word_bool_enabled(child):
            tag = "u"
        elif name in {"strike", "dstrike"} and word_bool_enabled(child):
            tag = "s"
        if tag and tag not in seen:
            open_parts.append(f"<{tag}>")
            close_parts.append(f"</{tag}>")
            seen.add(tag)
    return "".join(open_parts), "".join(reversed(close_parts))


def should_tokenize_inline_child(child: Any) -> bool:
    name = local_name(child)
    if is_math_element(child):
        return True
    if name == "hyperlink":
        return True
    if name == "r":
        if child.find(f".//{{{W_NS}}}drawing") is not None:
            return True
        return False
    return False


def is_format_run(child: Any) -> bool:
    if local_name(child) != "r":
        return False
    if child.find(f".//{{{W_NS}}}drawing") is not None:
        return False
    opening, closing = run_format_tags(child)
    return bool(opening and closing)


def inline_child_has_drawing(child: Any) -> bool:
    name = local_name(child)
    if name == "drawing":
        return True
    if name == "r":
        return child.find(f".//{{{W_NS}}}drawing") is not None
    return False


def rich_token_text_and_kind(child: Any) -> tuple[str, str]:
    text = element_text_with_math(child)
    if text:
        return text, "math" if is_math_element(child) else local_name(child)
    if inline_child_has_drawing(child):
        return "[drawing]", "drawing"
    return "", local_name(child)


def rich_tokenize_children(
    children: list[Any], item_id: str, start_seq: int = 1
) -> tuple[str, list[dict[str, Any]]]:
    parts: list[str] = []
    tokens: list[dict[str, Any]] = []
    seq = start_seq
    for child in children:
        if should_tokenize_inline_child(child):
            text, kind = rich_token_text_and_kind(child)
            if not text:
                continue
            token = f"[[RT:{item_id}:{seq:03d}]]"
            parts.append(token)
            tokens.append(
                {
                    "token": token,
                    "text": text,
                    "kind": kind,
                    "preserve": True,
                }
            )
            seq += 1
            continue
        if is_format_run(child):
            text = element_text_with_math(child)
            if text:
                opening, closing = run_format_tags(child)
                parts.append(f"{opening}{text}{closing}")
            continue
        text = element_text_with_math(child)
        if text:
            parts.append(text)
    return "".join(parts).strip(), tokens


def rich_text_for_paragraph(paragraph: Paragraph, item_id: str) -> tuple[str, list[dict[str, Any]]]:
    rich_text, tokens = rich_tokenize_children(list(paragraph._p.iterchildren()), item_id)
    return (rich_text or paragraph_text(paragraph), tokens)


def rich_text_for_cell(cell: Any, item_id: str) -> tuple[str, list[dict[str, Any]]]:
    lines: list[str] = []
    tokens: list[dict[str, Any]] = []
    seq = 1
    for child in cell._tc.iterchildren():
        if not isinstance(child, CT_P):
            continue
        text, para_tokens = rich_tokenize_children(list(child.iterchildren()), item_id, start_seq=seq)
        seq += len(para_tokens)
        if text:
            lines.append(text)
        tokens.extend(para_tokens)
    rich_text = "\n".join(lines).strip()
    return (rich_text or cell_text(cell), tokens)


def rich_xml_map_for_paragraph(paragraph: Paragraph, item_id: str) -> dict[str, Any]:
    _text, tokens = rich_tokenize_children(list(paragraph._p.iterchildren()), item_id)
    token_names = [token["token"] for token in tokens]
    rich_children = [
        child
        for child in paragraph._p.iterchildren()
        if should_tokenize_inline_child(child) and rich_token_text_and_kind(child)[0]
    ]
    return {token: deepcopy(child) for token, child in zip(token_names, rich_children)}


def rich_xml_map_for_cell(cell: Any, item_id: str) -> dict[str, Any]:
    mapping: dict[str, Any] = {}
    seq = 1
    for child in cell._tc.iterchildren():
        if not isinstance(child, CT_P):
            continue
        _text, tokens = rich_tokenize_children(list(child.iterchildren()), item_id, start_seq=seq)
        seq += len(tokens)
        token_names = [token["token"] for token in tokens]
        rich_children = [
            item
            for item in child.iterchildren()
            if should_tokenize_inline_child(item) and rich_token_text_and_kind(item)[0]
        ]
        mapping.update({token: deepcopy(item) for token, item in zip(token_names, rich_children)})
    return mapping


def is_translatable_text(text: str) -> bool:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if not cleaned:
        return False
    if re.fullmatch(r"[\d\s.,;:%()<>≤≥=+\-/–—\[\]]+", cleaned):
        return False
    if len(cleaned) <= 4 and re.fullmatch(r"[A-Za-z]?\d+[A-Za-z]?", cleaned):
        return False
    return bool(re.search(r"[A-Za-z\u4e00-\u9fff]", cleaned))


def media_entries(source_docx: Path) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(source_docx) as archive:
            names = [name for name in archive.namelist() if name.startswith("word/media/")]
    except zipfile.BadZipFile:
        names = []
    for idx, name in enumerate(sorted(names), start=1):
        entries.append(
            {
                "media_id": f"F{idx:03d}",
                "kind": "image",
                "source_ref": name,
                "filename": Path(name).name,
                "source_order": idx,
            }
        )
    return entries


def source_docx_quality_report(source_docx: Path) -> dict[str, Any]:
    ensure_docx_available()
    if not source_docx.exists():
        fail(f"source DOCX not found: {source_docx}")
    doc = Document(str(source_docx))
    texts = [paragraph_text(paragraph) for paragraph in doc.paragraphs]
    nonempty = [text for text in texts if text.strip()]
    table_count = len(doc.tables)
    media_count = len(media_entries(source_docx))
    space_aligned = [idx for idx, text in enumerate(nonempty, start=1) if re.search(r" {6,}", text)]
    very_long = [idx for idx, text in enumerate(nonempty, start=1) if len(text) >= 1200]
    page_artifacts = [
        idx
        for idx, text in enumerate(nonempty, start=1)
        if re.search(r"\bfrontiersin\.org\b|\bFrontiers in\b|\b10\.3389/", text, re.I)
    ]
    mixed_columns = [
        idx
        for idx, text in enumerate(nonempty, start=1)
        if (
            ("CITATION" in text and "Background:" in text)
            or ("Introduction" in text and "Method" in text and len(text) >= 400)
            or ("COPYRIGHT" in text and "Data," in text)
            or ("TABLE" in text and len(text) >= 300 and re.search(r" {6,}", text))
        )
    ]
    risk_reasons: list[str] = []
    if table_count == 0 and len(nonempty) >= 80 and len(space_aligned) >= 10:
        risk_reasons.append("many space-aligned paragraphs but no Word tables")
    if table_count == 0 and len(very_long) >= 3:
        risk_reasons.append("multiple very long paragraphs but no Word tables")
    if len(page_artifacts) >= 6:
        risk_reasons.append("repeated page header/footer artifacts detected")
    if len(mixed_columns) >= 2:
        risk_reasons.append("likely double-column text merged into single paragraphs")

    status = "blocked" if risk_reasons else "passed"
    return {
        "schema_version": SCHEMA_VERSION,
        "source_docx": str(source_docx),
        "status": status,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_count": len(nonempty),
        "table_count": table_count,
        "media_count": media_count,
        "space_aligned_paragraph_count": len(space_aligned),
        "very_long_paragraph_count": len(very_long),
        "page_artifact_paragraph_count": len(page_artifacts),
        "mixed_column_paragraph_count": len(mixed_columns),
        "sample_problem_paragraphs": sorted(set(space_aligned[:3] + very_long[:3] + mixed_columns[:3])),
        "risk_reasons": risk_reasons,
        "recommendation": (
            "Regenerate the DOCX from the original PDF with the MCP/OCR document conversion path, "
            "then rerun prepare-translation. Use --allow-layout-risk only for deliberate physical-layout drafts."
            if risk_reasons
            else "Source DOCX structure passed the basic layout-risk preflight."
        ),
    }


def preflight_source(bundle: Path, source_docx: Path | None = None, allow_layout_risk: bool = False) -> dict[str, Any]:
    ensure_bundle_dirs(bundle)
    path = source_docx or (bundle / "source.docx")
    report = source_docx_quality_report(path)
    errors: list[str] = []
    if report["status"] == "blocked" and not allow_layout_risk:
        errors.append(
            "source DOCX looks like a physical PDF layout extraction; regenerate it through the MCP/OCR conversion path before translation"
        )
    report["error_count"] = len(errors)
    report["errors"] = errors
    write_json(bundle / "logs" / "source_quality_report.json", report)
    return report


def init_bundle(bundle: Path, source_docx: Path | None = None, source_input: str = "") -> dict[str, Any]:
    ensure_bundle_dirs(bundle)
    if source_docx:
        if not source_docx.exists():
            fail(f"source DOCX not found: {source_docx}")
        target = bundle / "source.docx"
        if source_docx.resolve() != target.resolve():
            shutil.copy2(source_docx, target)
    manifest = {
        "schema_version": SCHEMA_VERSION,
        "created_at": now_stamp(),
        "workflow": "chatgpt-web-json-reader",
        "source_input": source_input,
        "source_docx": "source.docx" if (bundle / "source.docx").exists() else "",
        "stages": {
            "translation": "pending",
            "qa": "pending",
            "render": "pending",
            "gate": "pending",
        },
    }
    write_json(bundle / "bundle_manifest.json", manifest)
    return manifest


def guess_title(items: list[dict[str, Any]]) -> str:
    for item in items:
        text = str(item.get("original_text") or "").strip()
        if item.get("item_type") == "heading" and 8 <= len(text) <= 240:
            return text
    for item in items:
        text = str(item.get("original_text") or "").strip()
        if 20 <= len(text) <= 240:
            return text
    return "ChatGPT Web JSON Reader"


def table_item_sha256(cells: list[dict[str, Any]]) -> str:
    parts = [
        f"{cell.get('cell_id')}:{cell.get('original_sha256')}"
        for cell in cells
        if isinstance(cell, dict)
    ]
    return sha256_text("\n".join(parts))


def table_item_char_count(cells: list[dict[str, Any]]) -> int:
    return sum(len(str(cell.get("translation_source_text") or cell.get("original_text") or "")) for cell in cells if isinstance(cell, dict))


def public_table_cells(cells: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "cell_id": cell["cell_id"],
            "row": cell["row"],
            "col": cell["col"],
            "original_text": cell["original_text"],
            "translation_source_text": cell.get("translation_source_text", cell["original_text"]),
            "inline_tokens": cell.get("inline_tokens", []),
            "original_sha256": cell["original_sha256"],
        }
        for cell in cells
        if isinstance(cell, dict) and cell.get("translatable")
    ]


def build_ir(bundle: Path, title: str = "") -> dict[str, Any]:
    ensure_docx_available()
    ensure_bundle_dirs(bundle)
    source_docx = bundle / "source.docx"
    if not source_docx.exists():
        fail(f"missing source.docx in bundle: {source_docx}")

    document_id = f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{sha256_bytes(source_docx)[:8]}"
    doc = Document(str(source_docx))
    blocks: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    items: list[dict[str, Any]] = []
    current_section = ""
    paragraph_item_count = 0

    for entry in iter_docx_blocks(doc):
        if entry["kind"] == "paragraph":
            paragraph = entry["object"]
            text = paragraph_text(paragraph)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            item_type, heading_confidence = classify_paragraph_info(text, style_name)
            if item_type == "heading":
                current_section = text
            paragraph_item_count += 1
            item_id = f"S{paragraph_item_count:03d}"
            translation_source_text, inline_tokens = rich_text_for_paragraph(paragraph, item_id)
            item = {
                "item_id": item_id,
                "item_type": item_type,
                "location": {
                    "kind": "paragraph",
                    "docx_block_index": entry["docx_block_index"],
                    "paragraph_index": entry["paragraph_index"],
                },
                "section_path": [current_section] if current_section else [],
                "order": len(items) + 1,
                "original_text": text,
                "translation_source_text": translation_source_text,
                "inline_tokens": inline_tokens,
                "original_sha256": sha256_text(text),
                "style": style_name,
                "heading_confidence": heading_confidence,
                "char_count": len(text),
                "word_count": count_words(text),
            }
            items.append(item)
            blocks.append(item.copy())
            continue

        table = entry["object"]
        table_id = f"T{len(tables) + 1:03d}"
        table_record = {
            "table_id": table_id,
            "docx_block_index": entry["docx_block_index"],
            "table_index": entry["table_index"],
            "rows": len(table.rows),
            "columns": max((len(row.cells) for row in table.rows), default=0),
            "cells": [],
            "merged_aliases": [],
        }
        seen_table_cells: dict[str, str] = {}
        translatable_cells: list[dict[str, Any]] = []
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                key = cell_key(cell)
                if key in seen_table_cells:
                    table_record["merged_aliases"].append(
                        {
                            "row": r_idx,
                            "col": c_idx,
                            "canonical_cell_id": seen_table_cells[key],
                        }
                    )
                    continue
                text = cell_text(cell)
                if not text:
                    continue
                cell_id = f"{table_id}_R{r_idx:03d}_C{c_idx:03d}"
                translation_source_text, inline_tokens = rich_text_for_cell(cell, cell_id)
                seen_table_cells[key] = cell_id
                cell_record = {
                    "cell_id": cell_id,
                    "row": r_idx,
                    "col": c_idx,
                    "original_text": text,
                    "translation_source_text": translation_source_text,
                    "inline_tokens": inline_tokens,
                    "original_sha256": sha256_text(text),
                    "translatable": is_translatable_text(text),
                }
                table_record["cells"].append(cell_record)
                if cell_record["translatable"]:
                    translatable_cells.append(cell_record)
        tables.append(table_record)
        model_cells = public_table_cells(translatable_cells)
        if model_cells:
            table_item = {
                "item_id": table_id,
                "item_type": "table",
                "location": {
                    "kind": "table",
                    "docx_block_index": entry["docx_block_index"],
                    "table_index": entry["table_index"],
                    "table_id": table_id,
                },
                "section_path": [current_section] if current_section else [],
                "order": len(items) + 1,
                "original_text": f"Table {table_id} ({table_record['rows']} rows x {table_record['columns']} columns; {len(model_cells)} translatable cells)",
                "translation_source_text": "",
                "original_sha256": table_item_sha256(model_cells),
                "rows": table_record["rows"],
                "columns": table_record["columns"],
                "translatable_cell_count": len(model_cells),
                "skipped_cell_count": len(table_record["cells"]) - len(model_cells),
                "cells": model_cells,
                "char_count": table_item_char_count(model_cells),
                "word_count": sum(count_words(str(cell.get("original_text") or "")) for cell in model_cells),
            }
            items.append(table_item)

    ir = {
        "schema_version": SCHEMA_VERSION,
        "document_id": document_id,
        "source": {
            "source_docx": "source.docx",
            "source_sha256": sha256_bytes(source_docx),
            "title": title or guess_title(items),
        },
        "items": items,
        "blocks": blocks,
        "tables": tables,
        "media": media_entries(source_docx),
    }
    write_json(bundle / "document_ir.json", ir)
    write_json(
        bundle / "source_map.json",
        {
            "schema_version": SCHEMA_VERSION,
            "document_id": document_id,
            "source_docx": "source.docx",
            "items": items,
        },
    )
    write_json(
        bundle / "media_map.json",
        {
            "schema_version": SCHEMA_VERSION,
            "document_id": document_id,
            "media": ir["media"],
            "tables": tables,
        },
    )
    report = {
        "document_id": document_id,
        "item_count": len(items),
        "paragraph_item_count": paragraph_item_count,
        "table_count": len(tables),
        "table_item_count": sum(1 for item in items if item.get("item_type") == "table"),
        "translatable_table_cell_count": sum(len(item.get("cells", [])) for item in items if item.get("item_type") == "table"),
        "table_cell_item_count": sum(1 for item in items if item.get("item_type") == "table_cell"),
        "media_count": len(ir["media"]),
    }
    write_json(bundle / "logs" / "build_ir_report.json", report)
    return report


def item_group_char_count(group: list[dict[str, Any]]) -> int:
    return sum(int(item.get("char_count") or len(str(item.get("original_text") or ""))) for item in group)


def group_exceeds_limits(group: list[dict[str, Any]], max_items: int = 0, max_chars: int = 0) -> bool:
    if max_items and len(group) > max_items:
        return True
    if max_chars and item_group_char_count(group) > max_chars:
        return True
    return False


def table_row(item: dict[str, Any]) -> int:
    loc = item.get("location", {}) if isinstance(item.get("location"), dict) else {}
    return int(loc.get("row") or 0)


def table_id_for_item(item: dict[str, Any]) -> str:
    loc = item.get("location", {}) if isinstance(item.get("location"), dict) else {}
    return str(loc.get("table_id") or "")


def table_row_range(group: list[dict[str, Any]]) -> tuple[int, int]:
    rows = [table_row(item) for item in group if table_row(item)]
    if not rows:
        return 0, 0
    return min(rows), max(rows)


def split_table_group_by_rows(
    table_group: list[dict[str, Any]], max_items: int = 0, max_chars: int = 0
) -> list[list[dict[str, Any]]]:
    row_groups: list[list[dict[str, Any]]] = []
    idx = 0
    while idx < len(table_group):
        row = table_row(table_group[idx])
        row_group: list[dict[str, Any]] = []
        while idx < len(table_group) and table_row(table_group[idx]) == row:
            row_group.append(table_group[idx])
            idx += 1
        row_groups.append(row_group)

    packed: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_chars = 0
    for row_group in row_groups:
        row_chars = item_group_char_count(row_group)
        would_exceed = bool(
            current
            and (
                (max_items and len(current) + len(row_group) > max_items)
                or (max_chars and current_chars + row_chars > max_chars)
            )
        )
        if would_exceed:
            packed.append(current)
            current = []
            current_chars = 0
        current.extend(row_group)
        current_chars += row_chars
    if current:
        packed.append(current)
    return packed


def semantic_item_groups(
    items: list[dict[str, Any]], max_items: int = 0, max_chars: int = 0
) -> list[list[dict[str, Any]]]:
    """Create atomic groups so headings and tables are split only at stable boundaries."""

    groups: list[list[dict[str, Any]]] = []
    idx = 0
    while idx < len(items):
        item = items[idx]
        item_type = item.get("item_type")
        loc = item.get("location", {}) if isinstance(item.get("location"), dict) else {}
        if item_type == "table":
            groups.append([item])
            idx += 1
            continue
        if item_type == "table_cell":
            table_id = loc.get("table_id")
            group: list[dict[str, Any]] = []
            while idx < len(items):
                candidate = items[idx]
                candidate_loc = candidate.get("location", {}) if isinstance(candidate.get("location"), dict) else {}
                if candidate.get("item_type") != "table_cell" or candidate_loc.get("table_id") != table_id:
                    break
                group.append(candidate)
                idx += 1
            if group_exceeds_limits(group, max_items=max_items, max_chars=max_chars):
                groups.extend(split_table_group_by_rows(group, max_items=max_items, max_chars=max_chars))
            else:
                groups.append(group)
            continue

        if item_type == "heading":
            group = [item]
            idx += 1
            group_chars = item_group_char_count(group)
            while idx < len(items) and items[idx].get("item_type") not in {"heading", "table", "table_cell"}:
                candidate = items[idx]
                candidate_chars = item_group_char_count([candidate])
                if group and (
                    (max_items and len(group) + 1 > max_items)
                    or (max_chars and group_chars + candidate_chars > max_chars)
                ):
                    break
                group.append(candidate)
                group_chars += candidate_chars
                idx += 1
            groups.append(group)
            continue

        groups.append([item])
        idx += 1
    return groups


def chunk_items(items: list[dict[str, Any]], max_items: int, max_chars: int) -> list[list[dict[str, Any]]]:
    chunks: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_chars = 0
    for group in semantic_item_groups(items, max_items=max_items, max_chars=max_chars):
        group_chars = item_group_char_count(group)
        if current and (len(current) + len(group) > max_items or current_chars + group_chars > max_chars):
            chunks.append(current)
            current = []
            current_chars = 0
        current.extend(group)
        current_chars += group_chars
    if current:
        chunks.append(current)
    return chunks


def balanced_translation_groups(
    items: list[dict[str, Any]], max_items: int, max_chars: int
) -> list[list[dict[str, Any]]]:
    """Use fine semantic groups for translation load balancing.

    The greedy section chunker keeps a heading with as much following prose as
    possible. That is safe but can produce 96/31 style splits. For translation
    jobs, protect only the heading plus its first paragraph, then let later
    paragraphs participate in contiguous load balancing.
    """

    groups: list[list[dict[str, Any]]] = []
    idx = 0
    while idx < len(items):
        item = items[idx]
        item_type = item.get("item_type")
        if item_type in {"table", "table_cell"}:
            groups.append([item])
            idx += 1
            continue
        if item_type == "heading":
            group = [item]
            idx += 1
            if idx < len(items) and items[idx].get("item_type") not in {"heading", "table", "table_cell"}:
                candidate = items[idx]
                candidate_group = group + [candidate]
                if not group_exceeds_limits(candidate_group, max_items=max_items, max_chars=max_chars):
                    group = candidate_group
                    idx += 1
            groups.append(group)
            continue
        groups.append([item])
        idx += 1
    return groups


def flatten_groups(groups: list[list[dict[str, Any]]]) -> list[dict[str, Any]]:
    return [item for group in groups for item in group]


def balanced_chunk_items(items: list[dict[str, Any]], max_items: int, max_chars: int) -> list[list[dict[str, Any]]]:
    groups = balanced_translation_groups(items, max_items=max_items, max_chars=max_chars)
    if not groups:
        return []

    group_item_counts = [len(group) for group in groups]
    group_char_counts = [item_group_char_count(group) for group in groups]
    total_items = sum(group_item_counts)
    total_chars = sum(group_char_counts)
    min_by_items = ((total_items + max_items - 1) // max_items) if max_items else 1
    min_by_chars = ((total_chars + max_chars - 1) // max_chars) if max_chars else 1
    min_chunks = max(1, min_by_items, min_by_chars)

    prefix_items = [0]
    prefix_chars = [0]
    for item_count, char_count in zip(group_item_counts, group_char_counts):
        prefix_items.append(prefix_items[-1] + item_count)
        prefix_chars.append(prefix_chars[-1] + char_count)

    def segment_counts(start: int, end: int) -> tuple[int, int]:
        return prefix_items[end] - prefix_items[start], prefix_chars[end] - prefix_chars[start]

    def segment_allowed(start: int, end: int) -> bool:
        item_count, char_count = segment_counts(start, end)
        if end - start == 1:
            return True
        return (not max_items or item_count <= max_items) and (not max_chars or char_count <= max_chars)

    n = len(groups)
    for chunk_count in range(min_chunks, n + 1):
        target_items = total_items / chunk_count if total_items else 1.0
        target_chars = total_chars / chunk_count if total_chars else 1.0
        dp: list[list[tuple[float, int] | None]] = [[None] * (n + 1) for _ in range(chunk_count + 1)]
        dp[0][0] = (0.0, -1)
        for used in range(1, chunk_count + 1):
            for end in range(used, n + 1):
                best: tuple[float, int] | None = None
                for start in range(used - 1, end):
                    prev = dp[used - 1][start]
                    if prev is None or not segment_allowed(start, end):
                        continue
                    item_count, char_count = segment_counts(start, end)
                    item_score = ((item_count - target_items) / target_items) ** 2 if target_items else 0.0
                    char_score = ((char_count - target_chars) / target_chars) ** 2 if target_chars else 0.0
                    # Characters are the main proxy for model load; item count
                    # still matters so a long tail of tiny paragraphs is not ignored.
                    score = prev[0] + (0.65 * char_score) + (0.35 * item_score)
                    if best is None or score < best[0]:
                        best = (score, start)
                dp[used][end] = best
        if dp[chunk_count][n] is None:
            continue
        ranges: list[tuple[int, int]] = []
        end = n
        used = chunk_count
        while used > 0:
            entry = dp[used][end]
            if entry is None:
                break
            start = entry[1]
            ranges.append((start, end))
            end = start
            used -= 1
        if len(ranges) == chunk_count:
            ranges.reverse()
            return [flatten_groups(groups[start:end]) for start, end in ranges]

    return chunk_items(items, max_items=max_items, max_chars=max_chars)


def text_matches_auxiliary_section(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", text).strip().lower()
    if not normalized:
        return False
    return any(pattern in normalized for pattern in AUXILIARY_SECTION_PATTERNS)


def translation_lane_for_item(item: dict[str, Any]) -> str:
    if item.get("item_type") in {"table", "table_cell"}:
        return "auxiliary"
    section_path = item.get("section_path", [])
    if isinstance(section_path, list) and any(text_matches_auxiliary_section(str(part)) for part in section_path):
        return "auxiliary"
    if item.get("item_type") == "heading" and text_matches_auxiliary_section(str(item.get("original_text") or "")):
        return "auxiliary"
    return "body"


def partition_translation_lanes(items: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    lanes: dict[str, list[dict[str, Any]]] = {"body": [], "auxiliary": []}
    for item in items:
        lanes[translation_lane_for_item(item)].append(item)
    return lanes


def make_translation_chunks(
    items: list[dict[str, Any]],
    body_max_items: int,
    body_max_chars: int,
    aux_max_items: int,
    aux_max_chars: int,
) -> list[dict[str, Any]]:
    lanes = partition_translation_lanes(items)
    planned: list[dict[str, Any]] = []
    for lane, lane_items, max_items, max_chars in (
        ("body", lanes["body"], body_max_items, body_max_chars),
        ("auxiliary", lanes["auxiliary"], aux_max_items, aux_max_chars),
    ):
        for chunk in balanced_chunk_items(lane_items, max_items=max_items, max_chars=max_chars):
            planned.append(
                {
                    "lane": lane,
                    "items": chunk,
                    "max_items": max_items,
                    "max_chars": max_chars,
                    "chunker": "balanced_contiguous",
                }
            )
    return planned


def oversize_group_warnings(
    items: list[dict[str, Any]], max_items: int, max_chars: int
) -> list[dict[str, Any]]:
    warnings: list[dict[str, Any]] = []
    idx = 0
    while idx < len(items):
        item = items[idx]
        if item.get("item_type") == "table":
            if group_exceeds_limits([item], max_items=max_items, max_chars=max_chars):
                warnings.append(
                    {
                        "type": "oversize_table_item",
                        "message": "A compact table item exceeds job limits and will be kept as an atomic table job.",
                        "table_id": table_id_for_item(item),
                        "item_count": 1,
                        "translatable_cell_count": len(item.get("cells", [])),
                        "char_count": item_group_char_count([item]),
                        "max_items": max_items,
                        "max_chars": max_chars,
                    }
                )
            idx += 1
            continue
        if item.get("item_type") != "table_cell":
            idx += 1
            continue
        table_id = table_id_for_item(item)
        table_group: list[dict[str, Any]] = []
        while idx < len(items) and items[idx].get("item_type") == "table_cell" and table_id_for_item(items[idx]) == table_id:
            table_group.append(items[idx])
            idx += 1
        if group_exceeds_limits(table_group, max_items=max_items, max_chars=max_chars):
            row_start, row_end = table_row_range(table_group)
            warnings.append(
                {
                    "type": "oversize_table_row_fallback",
                    "message": "A table exceeds job limits and will be split by row groups to keep each row atomic.",
                    "table_id": table_id,
                    "row_start": row_start,
                    "row_end": row_end,
                    "item_count": len(table_group),
                    "char_count": item_group_char_count(table_group),
                    "max_items": max_items,
                    "max_chars": max_chars,
                }
            )

    for group in semantic_item_groups(items, max_items=max_items, max_chars=max_chars):
        char_count = item_group_char_count(group)
        if len(group) <= max_items and char_count <= max_chars:
            continue
        first = group[0]
        loc = first.get("location", {}) if isinstance(first.get("location"), dict) else {}
        warnings.append(
            {
                "type": "oversize_semantic_group",
                "message": "A table or heading-bound group exceeds job limits and will stay in one job to preserve structure.",
                "item_count": len(group),
                "char_count": char_count,
                "max_items": max_items,
                "max_chars": max_chars,
                "first_item_id": first.get("item_id", ""),
                "last_item_id": group[-1].get("item_id", ""),
                "group_kind": "table_row_group" if first.get("item_type") == "table_cell" else str(first.get("item_type", "")),
                "table_id": loc.get("table_id", ""),
                "row_start": table_row_range(group)[0],
                "row_end": table_row_range(group)[1],
                "suggestion": "This usually means one table row is itself larger than the configured job limits; consider raising limits or manually simplifying that source row.",
            }
        )
    return warnings


def section_context_for_item(item: dict[str, Any]) -> dict[str, Any]:
    path = item.get("section_path", [])
    if not isinstance(path, list):
        path = []
    return {
        "path": path,
        "nearest_heading": path[-1] if path else "",
    }


def chunk_context_for_items(chunk: list[dict[str, Any]]) -> dict[str, Any]:
    table_segments: list[dict[str, Any]] = []
    by_table: dict[str, list[dict[str, Any]]] = {}
    for item in chunk:
        if item.get("item_type") == "table":
            table_segments.append(
                {
                    "table_id": table_id_for_item(item),
                    "row_start": 1,
                    "row_end": int(item.get("rows") or 0),
                    "item_count": 1,
                    "translatable_cell_count": len(item.get("cells", [])),
                }
            )
            continue
        if item.get("item_type") != "table_cell":
            continue
        table_id = table_id_for_item(item)
        by_table.setdefault(table_id, []).append(item)
    for table_id, group in by_table.items():
        row_start, row_end = table_row_range(group)
        table_segments.append(
            {
                "table_id": table_id,
                "row_start": row_start,
                "row_end": row_end,
                "item_count": len(group),
            }
        )
    return {
        "table_segments": table_segments,
        "contains_heading": any(item.get("item_type") == "heading" for item in chunk),
    }


def model_item_payload(item: dict[str, Any]) -> dict[str, Any]:
    base = {
        "item_id": item["item_id"],
        "item_type": item["item_type"],
        "section_path": item.get("section_path", []),
        "section_context": section_context_for_item(item),
        "heading_confidence": item.get("heading_confidence", "none"),
        "order": item.get("order"),
        "location": item.get("location", {}),
        "original_sha256": item["original_sha256"],
    }
    if item.get("item_type") == "table":
        return {
            **base,
            "table_id": item.get("item_id"),
            "rows": item.get("rows", 0),
            "columns": item.get("columns", 0),
            "translatable_cell_count": item.get("translatable_cell_count", len(item.get("cells", []))),
            "skipped_cell_count": item.get("skipped_cell_count", 0),
            "cells": item.get("cells", []),
        }
    return {
        **base,
        "original_text": item["original_text"],
        "translation_source_text": item.get("translation_source_text", item["original_text"]),
        "inline_tokens": item.get("inline_tokens", []),
    }


def translated_cells_by_id(item: dict[str, Any]) -> dict[str, dict[str, Any]]:
    cells = item.get("translated_cells", [])
    if not isinstance(cells, list):
        return {}
    return {
        str(cell.get("cell_id")): cell
        for cell in cells
        if isinstance(cell, dict) and cell.get("cell_id")
    }


def stage_dir(bundle: Path, stage: str) -> Path:
    if stage not in {TRANSLATION_STAGE, QA_STAGE}:
        fail(f"unknown stage: {stage}")
    path = bundle / "jobs" / stage
    (path / "downloads").mkdir(parents=True, exist_ok=True)
    (path / "manifests").mkdir(parents=True, exist_ok=True)
    return path


def prompt_for_translation(job: dict[str, Any], expected_output: str) -> str:
    template = read_prompt_template("translation_user_template.md")
    return (
        f"{template}\n\n"
        "Output delivery requirement:\n"
        f"- Create or download one JSON file named exactly: {expected_output}\n"
        "- The file content must be the JSON object itself, not Markdown, not a prose explanation.\n"
        "- If the UI also shows a chat message, keep it empty or only provide the same raw JSON.\n\n"
        f"Input job_id: {job['job_id']}\n"
        f"Input document_id: {job['document_id']}\n\n"
        "Use the attached or pasted input JSON as input.json."
    )


def prompt_for_qa(job: dict[str, Any], expected_output: str) -> str:
    template = read_prompt_template("qa_user_template.md")
    return (
        f"{template}\n\n"
        "Output delivery requirement:\n"
        f"- Create or download one JSON file named exactly: {expected_output}\n"
        "- The file content must be the JSON object itself, not Markdown, not a prose explanation.\n"
        "- If the UI also shows a chat message, keep it empty or only provide the same raw JSON.\n\n"
        f"Input job_id: {job['job_id']}\n"
        f"Input document_id: {job['document_id']}\n\n"
        "Use the attached or pasted input JSON as qa_job_input.json."
    )


def write_exchange(stage_path: Path, stage: str, job_id: str, input_path: Path, prompt_path: Path) -> Path:
    expected_name = f"{job_id}_output.json"
    exchange = {
        "schema_version": SCHEMA_VERSION,
        "exchange_id": f"exchange_{job_id}",
        "provider": "chatgpt_web",
        "mode": "manual_upload_download",
        "stage": stage,
        "input_json": input_path.name,
        "prompt_file": prompt_path.name,
        "expected_output": {
            "kind": "json_file",
            "filename_hint": expected_name,
            "schema": "translation_output.schema.json" if stage == TRANSLATION_STAGE else "qa_output.schema.json",
        },
        "manual_steps": [
            "Open ChatGPT Web.",
            f"Upload or paste {input_path.name}.",
            f"Paste {prompt_path.name}.",
            f"Create/download one JSON file named exactly {expected_name}.",
            "Save it into this stage's downloads folder.",
        ],
    }
    exchange_path = stage_path / "manifests" / f"{job_id}_exchange.json"
    write_json(exchange_path, exchange)
    return exchange_path


def make_translation_jobs(
    bundle: Path,
    max_items: int = DEFAULT_BODY_TRANSLATION_MAX_ITEMS,
    max_chars: int = DEFAULT_BODY_TRANSLATION_MAX_CHARS,
    aux_max_items: int = DEFAULT_AUX_TRANSLATION_MAX_ITEMS,
    aux_max_chars: int = DEFAULT_AUX_TRANSLATION_MAX_CHARS,
) -> dict[str, Any]:
    ir = load_json_obj(bundle / "document_ir.json")
    items = ir.get("items")
    if not isinstance(items, list):
        fail("document_ir.json has no items list")
    lanes = partition_translation_lanes(items)
    warnings = [
        {**warning, "lane": "body"}
        for warning in oversize_group_warnings(lanes["body"], max_items=max_items, max_chars=max_chars)
    ] + [
        {**warning, "lane": "auxiliary"}
        for warning in oversize_group_warnings(lanes["auxiliary"], max_items=aux_max_items, max_chars=aux_max_chars)
    ]
    chunk_plan = make_translation_chunks(
        items,
        body_max_items=max_items,
        body_max_chars=max_chars,
        aux_max_items=aux_max_items,
        aux_max_chars=aux_max_chars,
    )
    out_dir = stage_dir(bundle, TRANSLATION_STAGE)

    jobs: list[dict[str, Any]] = []
    for idx, planned_chunk in enumerate(chunk_plan, start=1):
        chunk = planned_chunk["items"]
        lane = str(planned_chunk["lane"])
        job_id = f"translation_job_{idx:03d}"
        expected_output = f"{job_id}_output.json"
        job_items = [model_item_payload(item) for item in chunk]
        job = {
            "schema_version": SCHEMA_VERSION,
            "job_id": job_id,
            "task_type": "translate_items",
            "translation_lane": lane,
            "document_id": ir["document_id"],
            "source_sha256": ir["source"]["source_sha256"],
            "output_schema": "translation_output.schema.json",
            "instructions": {
                "target_language": "zh-CN",
                "style": "faithful_academic_reader",
                "do_not_summarize": True,
                "return_json_only": True,
                "output_delivery": {
                    "kind": "downloadable_json_file",
                    "filename": expected_output,
                },
            },
            "items": job_items,
            "chunking": {
                **chunk_context_for_items(chunk),
                "lane": lane,
                "max_items": planned_chunk["max_items"],
                "max_chars": planned_chunk["max_chars"],
                "chunker": planned_chunk.get("chunker", "greedy"),
            },
        }
        input_path = out_dir / f"{job_id}_input.json"
        prompt_path = out_dir / f"{job_id}_prompt.md"
        write_json(input_path, job)
        write_text(prompt_path, prompt_for_translation(job, expected_output) + "\n")
        exchange_path = write_exchange(out_dir, TRANSLATION_STAGE, job_id, input_path, prompt_path)
        jobs.append(
            {
                "job_id": job_id,
                "input_json": rel(input_path, bundle),
                "prompt_file": rel(prompt_path, bundle),
                "exchange_json": rel(exchange_path, bundle),
                "expected_output": f"jobs/translation/downloads/{expected_output}",
                "lane": lane,
                "item_count": len(job_items),
                "char_count": item_group_char_count(chunk),
                "chunking": job["chunking"],
            }
        )

    manifest = {
        "schema_version": SCHEMA_VERSION,
        "stage": TRANSLATION_STAGE,
        "created_at": now_stamp(),
        "job_count": len(jobs),
        "max_items": max_items,
        "max_chars": max_chars,
        "chunking_policy": {
            "body": {"max_items": max_items, "max_chars": max_chars},
            "auxiliary": {"max_items": aux_max_items, "max_chars": aux_max_chars},
            "auxiliary_rule": "Tables plus references/appendix/supplementary/administrative sections are packed separately from body text.",
        },
        "lane_counts": {
            "body_items": len(lanes["body"]),
            "auxiliary_items": len(lanes["auxiliary"]),
            "body_jobs": sum(1 for job in jobs if job.get("lane") == "body"),
            "auxiliary_jobs": sum(1 for job in jobs if job.get("lane") == "auxiliary"),
        },
        "warnings": warnings,
        "jobs": jobs,
        "downloads_dir": "jobs/translation/downloads",
    }
    write_json(out_dir / "translation_manifest.json", manifest)
    write_manual_steps(bundle, TRANSLATION_STAGE, manifest)
    return manifest


def write_manual_steps(bundle: Path, stage: str, manifest: dict[str, Any]) -> None:
    stage_path = stage_dir(bundle, stage)
    lines = [
        f"# Manual {stage} exchange",
        "",
        f"Upload/paste files from this folder: {stage_path}",
        f"Save downloaded JSON files into: {stage_path / 'downloads'}",
        f"Internal exchange manifests are archived in: {stage_path / 'manifests'}",
        "Only upload/paste the matching prompt and input JSON files.",
        "",
        "Jobs:",
    ]
    for job in manifest.get("jobs", []):
        lines.extend(
            [
                "",
                f"- {job['job_id']}",
                f"  - lane: {job.get('lane', stage)}",
                f"  - prompt: {Path(job['prompt_file']).name}",
                f"  - input: {Path(job['input_json']).name}",
                f"  - expected download: {Path(job['expected_output']).name}",
            ]
        )
    next_cmd = (
        f"python {Path(__file__).resolve()} accept-downloads --bundle {bundle} --stage {stage}"
    )
    lines.extend(["", "After downloads are saved, run:", "", f"```bash\n{next_cmd}\n```", ""])
    write_text(stage_path / "MANUAL_STEPS.md", "\n".join(lines))


def open_stage_folders(bundle: Path, stage: str) -> None:
    path = stage_dir(bundle, stage)
    open_path(path)
    open_path(path / "downloads")


def prepare_translation(
    bundle: Path,
    source_docx: Path | None,
    source_input: str,
    title: str,
    max_items: int,
    max_chars: int,
    aux_max_items: int,
    aux_max_chars: int,
    do_open: bool,
    allow_layout_risk: bool = False,
) -> dict[str, Any]:
    init_bundle(bundle, source_docx, source_input)
    quality_report = preflight_source(bundle, allow_layout_risk=allow_layout_risk)
    if quality_report["error_count"]:
        fail("source DOCX failed quality preflight; see logs/source_quality_report.json")
    build_report = build_ir(bundle, title)
    manifest = make_translation_jobs(
        bundle,
        max_items=max_items,
        max_chars=max_chars,
        aux_max_items=aux_max_items,
        aux_max_chars=aux_max_chars,
    )
    if do_open:
        open_stage_folders(bundle, TRANSLATION_STAGE)
    return {"source_quality": quality_report, "build_ir": build_report, "translation_jobs": manifest}


def expected_jobs(bundle: Path, stage: str) -> list[dict[str, Any]]:
    manifest_name = "translation_manifest.json" if stage == TRANSLATION_STAGE else "qa_manifest.json"
    manifest_path = stage_dir(bundle, stage) / manifest_name
    if not manifest_path.exists():
        fail(f"missing {manifest_path}; create jobs first")
    manifest = load_json_obj(manifest_path)
    jobs = manifest.get("jobs", [])
    if not isinstance(jobs, list):
        fail(f"{manifest_path} has invalid jobs list")
    return [job for job in jobs if isinstance(job, dict)]


def find_download_for_job(stage_path: Path, job_id: str, explicit_files: list[Path]) -> Path | None:
    expected = f"{job_id}_output.json"
    candidates = []
    for path in explicit_files:
        if path.name == expected or job_id in path.name:
            candidates.append(path)
    candidates.extend(sorted((stage_path / "downloads").glob(expected)))
    candidates.extend(sorted((stage_path / "downloads").glob(f"*{job_id}*.json")))
    for path in candidates:
        if path.exists() and path.is_file():
            return path
    return None


def normalize_download(stage: str, input_job: dict[str, Any], output: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
    errors: list[str] = []
    if output.get("job_id") != input_job.get("job_id"):
        errors.append(f"job_id mismatch: expected {input_job.get('job_id')}, got {output.get('job_id')}")
    if output.get("document_id") != input_job.get("document_id"):
        errors.append("document_id mismatch")
    if output.get("status") != "complete":
        errors.append(f"status is not complete: {output.get('status')!r}")

    if stage == TRANSLATION_STAGE:
        items = output.get("items")
        if not isinstance(items, list):
            errors.append("items must be a list")
            items = []
        normalized_items = []
        for item in items:
            if isinstance(item, dict):
                normalized_items.append(item)
            else:
                errors.append("all items must be objects")
        output["items"] = normalized_items
    else:
        findings = output.get("findings")
        if not isinstance(findings, list):
            errors.append("findings must be a list")
            output["findings"] = []
    return output, errors


def inline_format_tag_warnings(text: str, context: str) -> list[str]:
    warnings: list[str] = []
    stack: list[str] = []
    for match in INLINE_FORMAT_TAG_PATTERN.finditer(text):
        tag = match.group(0)
        name = tag[2:-1] if tag.startswith("</") else tag[1:-1]
        if not tag.startswith("</"):
            stack.append(name)
            continue
        if stack and stack[-1] == name:
            stack.pop()
            continue
        if name in stack:
            warnings.append(f"{context}: nested HTML format tags close out of order: {tag}")
            stack.remove(name)
        else:
            warnings.append(f"{context}: stray closing HTML format tag: {tag}")
    if stack:
        warnings.append(f"{context}: unclosed HTML format tags: {', '.join(stack)}")
    return warnings


def inline_format_tag_count_warnings(source_text: str, translated_text: str, context: str) -> list[str]:
    warnings: list[str] = []
    for tag in ("b", "i", "u", "s", "sup", "sub"):
        source_count = source_text.count(f"<{tag}>")
        translated_count = translated_text.count(f"<{tag}>")
        if source_count != translated_count:
            warnings.append(
                f"{context}: HTML format tag <{tag}> count changed from {source_count} to {translated_count}"
            )
    return warnings


def accept_downloads(bundle: Path, stage: str, files: list[Path]) -> dict[str, Any]:
    stage_path = stage_dir(bundle, stage)
    accepted: list[dict[str, Any]] = []
    errors: list[str] = []
    for job in expected_jobs(bundle, stage):
        job_id = str(job["job_id"])
        input_path = bundle / str(job["input_json"])
        input_job = load_json_obj(input_path)
        download = find_download_for_job(stage_path, job_id, files)
        if not download:
            errors.append(f"{job_id}: no downloaded JSON found in {stage_path / 'downloads'}")
            continue
        try:
            data = load_json_obj(download)
        except Exception as exc:
            errors.append(f"{job_id}: cannot parse {download}: {exc}")
            continue
        raw_path = stage_path / f"{job_id}_output.raw.json"
        normalized_path = stage_path / f"{job_id}_output.json"
        shutil.copy2(download, raw_path)
        normalized, normalization_errors = normalize_download(stage, input_job, data)
        write_json(normalized_path, normalized)
        validation = validate_one_output(bundle, stage, job_id, emit_file=False)
        validation["accept_errors"] = normalization_errors
        write_json(stage_path / f"{job_id}_validation.json", validation)
        accepted.append(
            {
                "job_id": job_id,
                "download": rel(download, bundle),
                "raw_output": rel(raw_path, bundle),
                "normalized_output": rel(normalized_path, bundle),
                "validation_status": "passed" if not validation["errors"] and not normalization_errors else "failed",
            }
        )
        errors.extend(f"{job_id}: {err}" for err in normalization_errors)
    report = {
        "stage": stage,
        "accepted_count": len(accepted),
        "accepted": accepted,
        "error_count": len(errors),
        "errors": errors,
    }
    write_json(stage_path / f"{stage}_accept_report.json", report)
    return report


def validate_one_output(bundle: Path, stage: str, job_id: str, emit_file: bool = True) -> dict[str, Any]:
    stage_path = stage_dir(bundle, stage)
    input_job = load_json_obj(stage_path / f"{job_id}_input.json")
    output_path = stage_path / f"{job_id}_output.json"
    errors: list[str] = []
    warnings: list[str] = []
    if not output_path.exists():
        errors.append(f"missing output: {rel(output_path, bundle)}")
        report = {"job_id": job_id, "stage": stage, "errors": errors, "warnings": warnings}
        if emit_file:
            write_json(stage_path / f"{job_id}_validation.json", report)
        return report

    try:
        output = load_json_obj(output_path)
    except Exception as exc:
        errors.append(f"cannot parse output JSON: {exc}")
        report = {"job_id": job_id, "stage": stage, "errors": errors, "warnings": warnings}
        if emit_file:
            write_json(stage_path / f"{job_id}_validation.json", report)
        return report

    if output.get("job_id") != input_job.get("job_id"):
        errors.append("job_id mismatch")
    if output.get("document_id") != input_job.get("document_id"):
        errors.append("document_id mismatch")
    if output.get("status") != "complete":
        errors.append("status must be complete")

    if stage == TRANSLATION_STAGE:
        input_items = input_job.get("items", [])
        output_items = output.get("items", [])
        if not isinstance(output_items, list):
            errors.append("items must be a list")
            output_items = []
        expected = {str(item.get("item_id")): item for item in input_items if isinstance(item, dict)}
        seen: dict[str, int] = {}
        for item in output_items:
            if not isinstance(item, dict):
                errors.append("output item is not an object")
                continue
            item_id = str(item.get("item_id") or "")
            seen[item_id] = seen.get(item_id, 0) + 1
            source_item = expected.get(item_id)
            if not source_item:
                errors.append(f"unexpected item_id: {item_id}")
                continue
            if item.get("original_sha256") != source_item.get("original_sha256"):
                errors.append(f"{item_id}: original_sha256 mismatch")
            if source_item.get("item_type") == "table":
                translated_cells = item.get("translated_cells")
                if not isinstance(translated_cells, list):
                    errors.append(f"{item_id}: translated_cells must be a list")
                    translated_cells = []
                expected_cells = {
                    str(cell.get("cell_id")): cell
                    for cell in source_item.get("cells", [])
                    if isinstance(cell, dict) and cell.get("cell_id")
                }
                seen_cells: dict[str, int] = {}
                for cell in translated_cells:
                    if not isinstance(cell, dict):
                        errors.append(f"{item_id}: translated cell is not an object")
                        continue
                    cell_id = str(cell.get("cell_id") or "")
                    seen_cells[cell_id] = seen_cells.get(cell_id, 0) + 1
                    source_cell = expected_cells.get(cell_id)
                    if not source_cell:
                        errors.append(f"{item_id}: unexpected cell_id {cell_id}")
                        continue
                    if cell.get("original_sha256") != source_cell.get("original_sha256"):
                        errors.append(f"{item_id}/{cell_id}: original_sha256 mismatch")
                    translated_cell_text = cell.get("translated_text")
                    for token in source_cell.get("inline_tokens", []):
                        token_name = token.get("token") if isinstance(token, dict) else ""
                        if token_name and isinstance(translated_cell_text, str) and token_name not in translated_cell_text:
                            errors.append(f"{item_id}/{cell_id}: missing rich-text token {token_name}")
                    if isinstance(translated_cell_text, str):
                        warnings.extend(inline_format_tag_warnings(translated_cell_text, f"{item_id}/{cell_id}"))
                        warnings.extend(
                            inline_format_tag_count_warnings(
                                str(source_cell.get("translation_source_text") or ""),
                                translated_cell_text,
                                f"{item_id}/{cell_id}",
                            )
                        )
                    if translated_cell_text is None and not cell.get("error"):
                        errors.append(f"{item_id}/{cell_id}: translated_text is null without structured error")
                    if isinstance(translated_cell_text, str) and not translated_cell_text.strip() and not cell.get("error"):
                        errors.append(f"{item_id}/{cell_id}: translated_text is empty")
                missing_cells = sorted(set(expected_cells) - set(seen_cells))
                duplicate_cells = sorted(cell_id for cell_id, count in seen_cells.items() if count > 1)
                if missing_cells:
                    errors.append(f"{item_id}: missing cell_ids: {', '.join(missing_cells[:20])}")
                if duplicate_cells:
                    errors.append(f"{item_id}: duplicate cell_ids: {', '.join(duplicate_cells[:20])}")
                continue
            translated = item.get("translated_text")
            expected_tokens = [
                token.get("token")
                for token in source_item.get("inline_tokens", [])
                if isinstance(token, dict) and token.get("token")
            ]
            for token in expected_tokens:
                if isinstance(translated, str) and token not in translated:
                    errors.append(f"{item_id}: missing rich-text token {token}")
            if isinstance(translated, str):
                warnings.extend(inline_format_tag_warnings(translated, item_id))
                warnings.extend(
                    inline_format_tag_count_warnings(
                        str(source_item.get("translation_source_text") or ""),
                        translated,
                        item_id,
                    )
                )
            if translated is None and not item.get("error"):
                errors.append(f"{item_id}: translated_text is null without structured error")
            if isinstance(translated, str) and not translated.strip() and not item.get("error"):
                errors.append(f"{item_id}: translated_text is empty")
        missing = sorted(set(expected) - set(seen))
        duplicates = sorted(item_id for item_id, count in seen.items() if count > 1)
        if missing:
            errors.append(f"missing item_ids: {', '.join(missing[:20])}")
        if duplicates:
            errors.append(f"duplicate item_ids: {', '.join(duplicates[:20])}")
    else:
        findings = output.get("findings", [])
        if not isinstance(findings, list):
            errors.append("findings must be a list")
            findings = []
        valid_item_ids = {
            str(item.get("item_id"))
            for item in input_job.get("items", [])
            if isinstance(item, dict)
        }
        input_items_by_id = {
            str(item.get("item_id")): item
            for item in input_job.get("items", [])
            if isinstance(item, dict)
        }
        for idx, finding in enumerate(findings, start=1):
            if not isinstance(finding, dict):
                errors.append(f"finding {idx}: not an object")
                continue
            item_id = str(finding.get("item_id") or "")
            if item_id and item_id not in valid_item_ids:
                errors.append(f"finding {idx}: unknown item_id {item_id}")
            if finding.get("apply") is True and not str(finding.get("suggested_translation") or "").strip():
                errors.append(f"finding {idx}: apply=true requires suggested_translation")
            if finding.get("apply") is True:
                source_item = input_items_by_id.get(item_id, {})
                suggestion = str(finding.get("suggested_translation") or "")
                token_source: list[Any] = []
                if isinstance(source_item, dict) and source_item.get("item_type") == "table":
                    cell_id = str(finding.get("cell_id") or "")
                    if not cell_id:
                        errors.append(f"finding {idx}: table apply=true requires cell_id")
                    cell_map = {
                        str(cell.get("cell_id")): cell
                        for cell in source_item.get("cells", [])
                        if isinstance(cell, dict)
                    }
                    source_cell = cell_map.get(cell_id)
                    if source_cell is None and cell_id:
                        errors.append(f"finding {idx}: unknown cell_id {cell_id}")
                    token_source = source_cell.get("inline_tokens", []) if isinstance(source_cell, dict) else []
                elif isinstance(source_item, dict):
                    token_source = source_item.get("inline_tokens", [])
                for token in token_source:
                    token_name = token.get("token") if isinstance(token, dict) else ""
                    if token_name and token_name not in suggestion:
                        errors.append(f"finding {idx}: suggested_translation missing rich-text token {token_name}")
                warnings.extend(inline_format_tag_warnings(suggestion, f"finding {idx}"))

    report = {
        "job_id": job_id,
        "stage": stage,
        "status": "passed" if not errors else "failed",
        "error_count": len(errors),
        "errors": errors,
        "warning_count": len(warnings),
        "warnings": warnings,
    }
    if emit_file:
        write_json(stage_path / f"{job_id}_validation.json", report)
    return report


def validate_model_outputs(bundle: Path, stage: str) -> dict[str, Any]:
    reports = []
    errors: list[str] = []
    warnings: list[str] = []
    for job in expected_jobs(bundle, stage):
        report = validate_one_output(bundle, stage, str(job["job_id"]))
        reports.append(report)
        errors.extend(f"{report['job_id']}: {err}" for err in report.get("errors", []))
        warnings.extend(f"{report['job_id']}: {warn}" for warn in report.get("warnings", []))
    stage_path = stage_dir(bundle, stage)
    summary = {
        "stage": stage,
        "job_count": len(reports),
        "passed_count": sum(1 for report in reports if not report.get("errors")),
        "error_count": len(errors),
        "errors": errors,
        "warning_count": len(warnings),
        "warnings": warnings,
        "reports": reports,
    }
    write_json(stage_path / f"{stage}_validation_summary.json", summary)
    return summary


def translation_map_from_outputs(bundle: Path) -> dict[str, dict[str, Any]]:
    translations: dict[str, dict[str, Any]] = {}
    for job in expected_jobs(bundle, TRANSLATION_STAGE):
        job_id = str(job["job_id"])
        output = load_json_obj(stage_dir(bundle, TRANSLATION_STAGE) / f"{job_id}_output.json")
        for item in output.get("items", []) or []:
            if isinstance(item, dict):
                item_id = str(item.get("item_id") or "")
                if item.get("item_type") == "table":
                    translations[item_id] = {
                        "item_id": item_id,
                        "item_type": "table",
                        "translated_cells": item.get("translated_cells", []),
                        "notes": item.get("notes", []),
                        "confidence": item.get("confidence", ""),
                        "source_job_id": job_id,
                        "qa_hint": item.get("qa_hint", {}),
                        "error": item.get("error"),
                    }
                else:
                    translations[item_id] = {
                        "item_id": item_id,
                        "translated_text": item.get("translated_text"),
                        "notes": item.get("notes", []),
                        "confidence": item.get("confidence", ""),
                        "source_job_id": job_id,
                        "qa_hint": item.get("qa_hint", {}),
                        "error": item.get("error"),
                    }
    return translations


def merge_draft(bundle: Path) -> dict[str, Any]:
    validation = validate_model_outputs(bundle, TRANSLATION_STAGE)
    if validation["error_count"]:
        fail("translation outputs failed validation; fix before merge-draft")
    ir = load_json_obj(bundle / "document_ir.json")
    translations = translation_map_from_outputs(bundle)
    merged_items: list[dict[str, Any]] = []
    missing: list[str] = []
    for item in ir.get("items", []):
        if not isinstance(item, dict):
            continue
        item_id = str(item["item_id"])
        translated = translations.get(item_id)
        if not translated:
            missing.append(item_id)
            continue
        if item.get("item_type") == "table":
            merged_items.append(
                {
                    **item,
                    "translated_cells": translated.get("translated_cells", []),
                    "translation_source_job_id": translated.get("source_job_id"),
                    "qa_hint": translated.get("qa_hint", {}),
                    "translation_notes": translated.get("notes", []),
                }
            )
        else:
            merged_items.append(
                {
                    **item,
                    "translated_text": translated.get("translated_text"),
                    "translation_source_job_id": translated.get("source_job_id"),
                    "qa_hint": translated.get("qa_hint", {}),
                    "translation_notes": translated.get("notes", []),
                }
            )
    draft = {
        "schema_version": SCHEMA_VERSION,
        "document_id": ir["document_id"],
        "source": ir["source"],
        "status": "draft",
        "items": merged_items,
        "missing_item_ids": missing,
    }
    write_json(bundle / "assembled" / "draft_assembled.json", draft)
    write_markdown(bundle / "outputs" / "draft_text.md", draft)
    report = {
        "merged_count": len(merged_items),
        "missing_count": len(missing),
        "missing_item_ids": missing[:50],
        "draft": "assembled/draft_assembled.json",
        "draft_markdown": "outputs/draft_text.md",
    }
    write_json(bundle / "assembled" / "merge_draft_report.json", report)
    return report


def make_qa_jobs(
    bundle: Path,
    max_items: int = DEFAULT_QA_MAX_ITEMS,
    max_chars: int = DEFAULT_QA_MAX_CHARS,
    do_open: bool = True,
) -> dict[str, Any]:
    draft_path = bundle / "assembled" / "draft_assembled.json"
    if not draft_path.exists():
        fail("missing assembled/draft_assembled.json; run merge-draft first")
    draft = load_json_obj(draft_path)
    items = [item for item in draft.get("items", []) if isinstance(item, dict)]
    chunks = chunk_items(items, max_items=max_items, max_chars=max_chars)
    out_dir = stage_dir(bundle, QA_STAGE)
    jobs: list[dict[str, Any]] = []
    for idx, chunk in enumerate(chunks, start=1):
        job_id = f"qa_job_{idx:03d}"
        expected_output = f"{job_id}_output.json"
        job_items = []
        for item in chunk:
            payload_item = model_item_payload(item)
            if item.get("item_type") == "table":
                payload_item["translated_cells"] = item.get("translated_cells", [])
                payload_item["translated_cell_count"] = len(item.get("translated_cells", []))
            else:
                payload_item["translated_text"] = item.get("translated_text")
                payload_item["translated_sha256"] = sha256_text(str(item.get("translated_text") or ""))
            payload_item["qa_hint"] = item.get("qa_hint", {})
            job_items.append(payload_item)
        job = {
            "schema_version": SCHEMA_VERSION,
            "job_id": job_id,
            "task_type": "qa_translation",
            "document_id": draft["document_id"],
            "draft_status": draft.get("status", "draft"),
            "output_schema": "qa_output.schema.json",
            "instructions": {
                "review_scope": (
                    "faithfulness, omissions, terminology consistency, table-cell accuracy, "
                    "references/citations/units/math preservation, and Chinese readability"
                ),
                "return_json_only": True,
                "do_not_rewrite_whole_document": True,
                "output_delivery": {
                    "kind": "downloadable_json_file",
                    "filename": expected_output,
                },
            },
            "items": job_items,
            "chunking": chunk_context_for_items(chunk),
        }
        input_path = out_dir / f"{job_id}_input.json"
        prompt_path = out_dir / f"{job_id}_prompt.md"
        write_json(input_path, job)
        write_text(prompt_path, prompt_for_qa(job, expected_output) + "\n")
        exchange_path = write_exchange(out_dir, QA_STAGE, job_id, input_path, prompt_path)
        jobs.append(
            {
                "job_id": job_id,
                "input_json": rel(input_path, bundle),
                "prompt_file": rel(prompt_path, bundle),
                "exchange_json": rel(exchange_path, bundle),
                "expected_output": f"jobs/qa/downloads/{expected_output}",
                "item_count": len(job_items),
                "chunking": chunk_context_for_items(chunk),
            }
        )
    manifest = {
        "schema_version": SCHEMA_VERSION,
        "stage": QA_STAGE,
        "created_at": now_stamp(),
        "job_count": len(jobs),
        "max_items": max_items,
        "max_chars": max_chars,
        "jobs": jobs,
        "downloads_dir": "jobs/qa/downloads",
    }
    write_json(out_dir / "qa_manifest.json", manifest)
    write_manual_steps(bundle, QA_STAGE, manifest)
    if do_open:
        open_stage_folders(bundle, QA_STAGE)
    return manifest


def apply_qa(bundle: Path) -> dict[str, Any]:
    validation = validate_model_outputs(bundle, QA_STAGE)
    if validation["error_count"]:
        fail("QA outputs failed validation; fix before apply-qa")
    draft = load_json_obj(bundle / "assembled" / "draft_assembled.json")
    by_id = {str(item.get("item_id")): dict(item) for item in draft.get("items", []) if isinstance(item, dict)}
    applied: list[dict[str, Any]] = []
    advisory: list[dict[str, Any]] = []
    for job in expected_jobs(bundle, QA_STAGE):
        job_id = str(job["job_id"])
        output = load_json_obj(stage_dir(bundle, QA_STAGE) / f"{job_id}_output.json")
        for finding in output.get("findings", []) or []:
            if not isinstance(finding, dict):
                continue
            item_id = str(finding.get("item_id") or "")
            if finding.get("apply") is True and item_id in by_id:
                suggested = str(finding.get("suggested_translation") or "").strip()
                if suggested:
                    cell_id = str(finding.get("cell_id") or "")
                    if by_id[item_id].get("item_type") == "table" and cell_id:
                        cells = translated_cells_by_id(by_id[item_id])
                        if cell_id not in cells:
                            advisory.append({"job_id": job_id, "finding": finding})
                            continue
                        before = cells[cell_id].get("translated_text")
                        cells[cell_id]["translated_text"] = suggested
                        by_id[item_id]["translated_cells"] = list(cells.values())
                    else:
                        before = by_id[item_id].get("translated_text")
                        by_id[item_id]["translated_text"] = suggested
                    by_id[item_id].setdefault("qa_applied_findings", []).append(
                        {
                            "job_id": job_id,
                            "finding_id": finding.get("finding_id", ""),
                            "severity": finding.get("severity", ""),
                            "category": finding.get("category", ""),
                            "problem": finding.get("problem", ""),
                            "cell_id": cell_id,
                            "previous_translation_sha256": sha256_text(str(before or "")),
                        }
                    )
                    applied.append({"job_id": job_id, "item_id": item_id, "cell_id": cell_id, "finding": finding})
                    continue
            advisory.append({"job_id": job_id, "finding": finding})

    final_items = [by_id[str(item.get("item_id"))] for item in draft.get("items", []) if str(item.get("item_id")) in by_id]
    payload = {
        "schema_version": SCHEMA_VERSION,
        "document_id": draft["document_id"],
        "source": draft.get("source", {}),
        "status": "qa_applied",
        "items": final_items,
        "qa_summary": {
            "applied_count": len(applied),
            "advisory_count": len(advisory),
        },
        "advisory_findings": advisory,
    }
    write_json(bundle / "assembled" / "review_applied.json", payload)
    write_json(bundle / "assembled" / "final_payload.json", payload)
    report = {
        "applied_count": len(applied),
        "advisory_count": len(advisory),
        "review_applied": "assembled/review_applied.json",
        "final_payload": "assembled/final_payload.json",
    }
    write_json(bundle / "assembled" / "apply_qa_report.json", report)
    return report


def promote_draft_without_qa(bundle: Path) -> dict[str, Any]:
    draft_path = bundle / "assembled" / "draft_assembled.json"
    if not draft_path.exists():
        fail("missing assembled/draft_assembled.json; run merge-draft first")
    draft = load_json_obj(draft_path)
    payload = {
        **draft,
        "status": "qa_skipped_fast",
        "qa_summary": {
            "skipped": True,
            "reason": "fast mode requested; QA stage was not run",
            "applied_count": 0,
            "advisory_count": 0,
        },
        "advisory_findings": [],
    }
    write_json(bundle / "assembled" / "final_payload.json", payload)
    report = {
        "mode": "fast",
        "qa_skipped": True,
        "source_draft": "assembled/draft_assembled.json",
        "final_payload": "assembled/final_payload.json",
        "item_count": len([item for item in payload.get("items", []) if isinstance(item, dict)]),
    }
    write_json(bundle / "assembled" / "fast_finish_report.json", report)
    return report


def finish_fast(bundle: Path, formats: str = "md,docx,json", do_open: bool = True) -> dict[str, Any]:
    accept_report = accept_downloads(bundle, TRANSLATION_STAGE, [])
    if accept_report["error_count"]:
        fail("translation downloads failed acceptance; fix downloads before finish-fast")
    validation = validate_model_outputs(bundle, TRANSLATION_STAGE)
    if validation["error_count"]:
        fail("translation outputs failed validation; fix before finish-fast")
    merge_report = merge_draft(bundle)
    fast_report = promote_draft_without_qa(bundle)
    render_report = render_final(bundle, formats)
    gate_report = gate(bundle)
    if do_open:
        open_path(bundle / "outputs")
    return {
        "mode": "fast",
        "qa_skipped": True,
        "accept_downloads": accept_report,
        "translation_validation": validation,
        "merge_draft": merge_report,
        "fast_finish": fast_report,
        "render_final": render_report,
        "gate": gate_report,
    }


def write_markdown(path: Path, payload: dict[str, Any]) -> None:
    lines = [
        f"# {payload.get('source', {}).get('title', 'Translated Reader')}",
        "",
        f"Document ID: `{payload.get('document_id', '')}`",
        "",
    ]
    for item in payload.get("items", []):
        if not isinstance(item, dict):
            continue
        original = str(item.get("original_text") or "").strip()
        translated = str(item.get("translated_text") or "").strip()
        if not original and not translated:
            continue
        translated = plain_text_from_rich_tokens(translated, item)
        item_type = item.get("item_type", "")
        item_id = item.get("item_id", "")
        if item_type == "heading":
            lines.extend([f"## {original}", "", translated, ""])
        elif item_type == "table":
            translated_cells = translated_cells_by_id(item)
            lines.extend([f"**{item_id}**", "", original, ""])
            for cell in item.get("cells", []):
                if not isinstance(cell, dict):
                    continue
                cell_id = str(cell.get("cell_id") or "")
                translated_cell = str(translated_cells.get(cell_id, {}).get("translated_text") or "").strip()
                if translated_cell:
                    translated_cell = plain_text_from_rich_tokens(translated_cell, cell)
                    lines.extend([f"- {cell_id} R{cell.get('row')}C{cell.get('col')}: {cell.get('original_text', '')}", f"  {translated_cell}"])
            lines.append("")
        elif item_type == "table_cell":
            lines.extend([f"**{item_id}**", "", original, "", translated, ""])
        else:
            lines.extend([original, "", translated, ""])
    write_text(path, "\n".join(lines).rstrip() + "\n")


def plain_text_from_rich_tokens(text: str, item: dict[str, Any]) -> str:
    token_map = {
        str(token.get("token")): str(token.get("text") or "")
        for token in item.get("inline_tokens", [])
        if isinstance(token, dict)
    }
    for token, replacement in token_map.items():
        text = text.replace(token, replacement)
    replacements = {
        "<b>": "**",
        "</b>": "**",
        "<i>": "*",
        "</i>": "*",
        "<u>": "",
        "</u>": "",
        "<s>": "~~",
        "</s>": "~~",
        "<sup>": "^(",
        "</sup>": ")",
        "<sub>": "_(",
        "</sub>": ")",
    }
    for tag, replacement in replacements.items():
        text = text.replace(tag, replacement)
    return text


def add_formatted_run(paragraph: Paragraph, text: str, active_formats: dict[str, int], size: float) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    if Pt is not None:
        run.font.size = Pt(size)
    if RGBColor is not None:
        run.font.color.rgb = RGBColor(31, 78, 121)
    if active_formats.get("b", 0) > 0:
        run.font.bold = True
    if active_formats.get("i", 0) > 0:
        run.font.italic = True
    if active_formats.get("u", 0) > 0:
        run.font.underline = True
    if active_formats.get("s", 0) > 0:
        run.font.strike = True
    if active_formats.get("sup", 0) > 0:
        run.font.superscript = True
    if active_formats.get("sub", 0) > 0:
        run.font.subscript = True


def add_blue_run(paragraph: Paragraph, text: str, size: float) -> None:
    add_formatted_run(paragraph, text, {}, size)


def add_rich_text(paragraph: Paragraph, text: str, token_xml_map: dict[str, Any], size: float) -> dict[str, int]:
    inserted_tokens = 0
    missing_tokens = 0
    active_formats: dict[str, int] = {}
    pos = 0
    for match in RICH_TOKEN_OR_FORMAT_TAG_PATTERN.finditer(text):
        add_formatted_run(paragraph, text[pos : match.start()], active_formats, size)
        marker = match.group(0)
        if marker.startswith("[[RT:"):
            xml = token_xml_map.get(marker)
            if xml is None:
                add_formatted_run(paragraph, marker, active_formats, size)
                missing_tokens += 1
            else:
                paragraph._p.append(deepcopy(xml))
                inserted_tokens += 1
        elif marker.startswith("</"):
            tag = marker[2:-1]
            if active_formats.get(tag, 0) > 1:
                active_formats[tag] -= 1
            else:
                active_formats.pop(tag, None)
        else:
            tag = marker[1:-1]
            active_formats[tag] = active_formats.get(tag, 0) + 1
        pos = match.end()
    add_formatted_run(paragraph, text[pos:], active_formats, size)
    return {"inserted_rich_tokens": inserted_tokens, "missing_rich_tokens": missing_tokens}


def insert_paragraph_after(paragraph: Paragraph, text: str, item_id: str) -> dict[str, int]:
    if OxmlElement is None:
        fail("python-docx OxmlElement is required to insert translated paragraphs")
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    try:
        inserted.style = paragraph.style
    except Exception:
        pass
    return add_rich_text(inserted, text, rich_xml_map_for_paragraph(paragraph, item_id), 10.5)


def append_cell_translation(cell: Any, text: str, item_id: str) -> dict[str, int] | None:
    text = str(text or "").strip()
    if not text or text in cell.text:
        return None
    para = cell.add_paragraph()
    return add_rich_text(para, text, rich_xml_map_for_cell(cell, item_id), 9)


def render_docx(bundle: Path, payload: dict[str, Any], output: Path) -> dict[str, Any]:
    ensure_docx_available()
    ir = load_json_obj(bundle / "document_ir.json")
    source = bundle / "source.docx"
    if not source.exists():
        fail("missing source.docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)
    doc = Document(str(output))
    by_id = {str(item.get("item_id")): item for item in payload.get("items", []) if isinstance(item, dict)}

    paragraph_translations: dict[int, tuple[str, str]] = {}
    cell_translations: dict[tuple[int, int, int], tuple[str, str]] = {}
    cell_id_locations: dict[str, tuple[int, int, int]] = {}
    for table_record in ir.get("tables", []):
        if not isinstance(table_record, dict):
            continue
        table_index = int(table_record.get("table_index") or 0)
        for cell_record in table_record.get("cells", []):
            if not isinstance(cell_record, dict):
                continue
            cell_id = str(cell_record.get("cell_id") or "")
            if cell_id:
                cell_id_locations[cell_id] = (
                    table_index,
                    int(cell_record.get("row") or 0),
                    int(cell_record.get("col") or 0),
                )
    for item in ir.get("items", []):
        if not isinstance(item, dict):
            continue
        rendered_item = by_id.get(str(item.get("item_id")), {})
        loc = item.get("location", {})
        if item.get("item_type") == "table":
            for translated_cell in rendered_item.get("translated_cells", []) if isinstance(rendered_item, dict) else []:
                if not isinstance(translated_cell, dict):
                    continue
                translated = str(translated_cell.get("translated_text") or "").strip()
                cell_id = str(translated_cell.get("cell_id") or "")
                key = cell_id_locations.get(cell_id)
                if translated and key:
                    cell_translations[key] = (cell_id, translated)
            continue
        translated = str(rendered_item.get("translated_text") or "").strip() if isinstance(rendered_item, dict) else ""
        if not translated:
            continue
        if loc.get("kind") == "paragraph":
            paragraph_translations[int(loc["paragraph_index"])] = (str(item.get("item_id")), translated)
        elif loc.get("kind") == "table_cell":
            key = (int(loc["table_index"]), int(loc["row"]), int(loc["col"]))
            cell_translations[key] = (str(item.get("item_id")), translated)

    inserted_paragraphs = 0
    inserted_cells = 0
    inserted_rich_tokens = 0
    missing_rich_tokens = 0
    # Insert from the end of the document toward the beginning so paragraph indexes
    # captured in document_ir.json remain stable while new translation paragraphs appear.
    for entry in reversed(list(iter_docx_blocks(doc))):
        if entry["kind"] == "paragraph":
            translation = paragraph_translations.get(int(entry["paragraph_index"]))
            if translation:
                token_report = insert_paragraph_after(entry["object"], translation[1], translation[0])
                inserted_rich_tokens += token_report["inserted_rich_tokens"]
                missing_rich_tokens += token_report["missing_rich_tokens"]
                inserted_paragraphs += 1
        elif entry["kind"] == "table":
            table = entry["object"]
            table_index = int(entry["table_index"])
            seen_cells: set[str] = set()
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    key = cell_key(cell)
                    if key in seen_cells:
                        continue
                    seen_cells.add(key)
                    translation = cell_translations.get((table_index, r_idx, c_idx))
                    token_report = append_cell_translation(cell, translation[1], translation[0]) if translation else None
                    if token_report is not None:
                        inserted_rich_tokens += token_report["inserted_rich_tokens"]
                        missing_rich_tokens += token_report["missing_rich_tokens"]
                        inserted_cells += 1

    # Post-process simple text LaTeX snippets so common statistics symbols are
    # readable in the final Word file without replacing native Office Math XML.
    latex_symbol_replacements = 0
    for p in doc.paragraphs:
        latex_symbol_replacements += _replace_math_symbols(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    latex_symbol_replacements += _replace_math_symbols(p)

    doc.save(output)
    return {
        "output": rel(output, bundle),
        "inserted_paragraph_translations": inserted_paragraphs,
        "inserted_table_cell_translations": inserted_cells,
        "inserted_rich_tokens": inserted_rich_tokens,
        "missing_rich_tokens": missing_rich_tokens,
        "latex_symbol_replacements": latex_symbol_replacements,
    }


def render_final(bundle: Path, formats: str) -> dict[str, Any]:
    payload_path = bundle / "assembled" / "final_payload.json"
    if not payload_path.exists():
        payload_path = bundle / "assembled" / "draft_assembled.json"
    if not payload_path.exists():
        fail("missing final_payload.json or draft_assembled.json")
    payload = load_json_obj(payload_path)
    outputs = bundle / "outputs"
    outputs.mkdir(exist_ok=True)
    requested = {part.strip() for part in formats.split(",") if part.strip()}
    report: dict[str, Any] = {"formats": sorted(requested), "outputs": {}}
    if "json" in requested:
        out_json = outputs / "final_payload.json"
        write_json(out_json, payload)
        report["outputs"]["json"] = rel(out_json, bundle)
    if "md" in requested:
        out_md = outputs / "final_text.md"
        write_markdown(out_md, payload)
        report["outputs"]["md"] = rel(out_md, bundle)
    if "docx" in requested:
        out_docx = outputs / "final_text.docx"
        report["outputs"]["docx"] = render_docx(bundle, payload, out_docx)
    write_json(outputs / "render_report.json", report)
    return report


def gate(bundle: Path) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    required = [
        "source.docx",
        "document_ir.json",
        "jobs/translation/translation_manifest.json",
        "assembled/draft_assembled.json",
        "outputs/final_text.md",
        "outputs/final_payload.json",
    ]
    for item in required:
        if not (bundle / item).exists():
            errors.append(f"missing {item}")
    translation_summary = bundle / "jobs/translation/translation_validation_summary.json"
    if translation_summary.exists():
        data = load_json_obj(translation_summary)
        if data.get("error_count"):
            errors.append("translation validation summary has errors")
    else:
        errors.append("missing translation validation summary")

    fast_report = bundle / "assembled/fast_finish_report.json"
    fast_mode = False
    if fast_report.exists():
        try:
            fast_mode = bool(load_json_obj(fast_report).get("qa_skipped"))
        except Exception:
            fast_mode = False

    qa_manifest = bundle / "jobs/qa/qa_manifest.json"
    if fast_mode:
        warnings.append("QA stage skipped by finish-fast")
    elif qa_manifest.exists():
        qa_summary = bundle / "jobs/qa/qa_validation_summary.json"
        if qa_summary.exists():
            data = load_json_obj(qa_summary)
            if data.get("error_count"):
                errors.append("QA validation summary has errors")
        else:
            errors.append("QA jobs exist but qa_validation_summary.json is missing")
        if not (bundle / "assembled/final_payload.json").exists():
            errors.append("QA jobs exist but final_payload.json is missing")
    else:
        warnings.append("QA stage has not been created")

    report = {
        "schema_version": SCHEMA_VERSION,
        "bundle": str(bundle),
        "gate_status": "passed" if not errors else "blocked",
        "error_count": len(errors),
        "errors": errors,
        "warnings": warnings,
    }
    write_json(bundle / "outputs" / "validation_report.json", report)
    return report


def status(bundle: Path) -> dict[str, Any]:
    checks = {
        "source.docx": (bundle / "source.docx").exists(),
        "document_ir.json": (bundle / "document_ir.json").exists(),
        "translation_manifest": (bundle / "jobs/translation/translation_manifest.json").exists(),
        "draft_assembled": (bundle / "assembled/draft_assembled.json").exists(),
        "qa_manifest": (bundle / "jobs/qa/qa_manifest.json").exists(),
        "final_payload": (bundle / "assembled/final_payload.json").exists(),
        "final_md": (bundle / "outputs/final_text.md").exists(),
        "final_docx": (bundle / "outputs/final_text.docx").exists(),
        "validation_report": (bundle / "outputs/validation_report.json").exists(),
    }
    return {"bundle": str(bundle), "checks": checks}


def make_repair_prompt(bundle: Path, stage: str, job_id: str) -> dict[str, Any]:
    stage_path = stage_dir(bundle, stage)
    validation_path = stage_path / f"{job_id}_validation.json"
    if not validation_path.exists():
        validate_one_output(bundle, stage, job_id)
    validation = load_json_obj(validation_path)
    input_text = (stage_path / f"{job_id}_input.json").read_text(encoding="utf-8")
    output_path = stage_path / f"{job_id}_output.raw.json"
    output_text = output_path.read_text(encoding="utf-8") if output_path.exists() else ""
    repair = read_prompt_template("repair_user_template.md")
    prompt = (
        f"{repair}\n\n"
        f"Stage: {stage}\n"
        f"Job ID: {job_id}\n"
        f"Validation errors:\n{json.dumps(validation.get('errors', []), ensure_ascii=False, indent=2)}\n\n"
        f"Original input JSON:\n{input_text}\n\n"
        f"Previous output JSON:\n{output_text}\n"
    )
    repair_path = stage_path / f"{job_id}_repair_prompt.md"
    write_text(repair_path, prompt)
    open_path(stage_path)
    return {"repair_prompt": rel(repair_path, bundle)}


def print_result(data: Any, as_json: bool = True) -> None:
    if as_json:
        print(json.dumps(data, ensure_ascii=False, indent=2))
    else:
        print(data)


def main() -> int:
    parser = argparse.ArgumentParser(description="ChatGPT Web JSON reader pipeline")
    sub = parser.add_subparsers(dest="command", required=True)

    def add_bundle(p: argparse.ArgumentParser) -> None:
        p.add_argument("--bundle", type=Path, required=True)

    p = sub.add_parser("init-bundle")
    add_bundle(p)
    p.add_argument("--source-docx", type=Path)
    p.add_argument("--source-input", default="")

    p = sub.add_parser("build-ir")
    add_bundle(p)
    p.add_argument("--title", default="")
    p.add_argument("--allow-layout-risk", action="store_true")

    p = sub.add_parser("make-translation-jobs")
    add_bundle(p)
    p.add_argument("--max-items", type=int, default=DEFAULT_BODY_TRANSLATION_MAX_ITEMS)
    p.add_argument("--max-chars", type=int, default=DEFAULT_BODY_TRANSLATION_MAX_CHARS)
    p.add_argument("--aux-max-items", type=int, default=DEFAULT_AUX_TRANSLATION_MAX_ITEMS)
    p.add_argument("--aux-max-chars", type=int, default=DEFAULT_AUX_TRANSLATION_MAX_CHARS)
    p.add_argument("--no-open", action="store_true")
    p.add_argument("--allow-layout-risk", action="store_true")

    p = sub.add_parser("preflight-source")
    add_bundle(p)
    p.add_argument("--source-docx", type=Path)
    p.add_argument("--allow-layout-risk", action="store_true")

    p = sub.add_parser("prepare-translation")
    add_bundle(p)
    p.add_argument("--source-docx", type=Path, required=True)
    p.add_argument("--source-input", default="")
    p.add_argument("--title", default="")
    p.add_argument("--max-items", type=int, default=DEFAULT_BODY_TRANSLATION_MAX_ITEMS)
    p.add_argument("--max-chars", type=int, default=DEFAULT_BODY_TRANSLATION_MAX_CHARS)
    p.add_argument("--aux-max-items", type=int, default=DEFAULT_AUX_TRANSLATION_MAX_ITEMS)
    p.add_argument("--aux-max-chars", type=int, default=DEFAULT_AUX_TRANSLATION_MAX_CHARS)
    p.add_argument("--no-open", action="store_true")
    p.add_argument("--allow-layout-risk", action="store_true")

    p = sub.add_parser("open-hand-off")
    add_bundle(p)
    p.add_argument("--stage", choices=[TRANSLATION_STAGE, QA_STAGE], required=True)

    p = sub.add_parser("accept-downloads")
    add_bundle(p)
    p.add_argument("--stage", choices=[TRANSLATION_STAGE, QA_STAGE], required=True)
    p.add_argument("--file", type=Path, action="append", default=[])

    p = sub.add_parser("validate-model-outputs")
    add_bundle(p)
    p.add_argument("--stage", choices=[TRANSLATION_STAGE, QA_STAGE], required=True)

    p = sub.add_parser("merge-draft")
    add_bundle(p)

    p = sub.add_parser("make-qa-jobs")
    add_bundle(p)
    p.add_argument("--max-items", type=int, default=DEFAULT_QA_MAX_ITEMS)
    p.add_argument("--max-chars", type=int, default=DEFAULT_QA_MAX_CHARS)
    p.add_argument("--no-open", action="store_true")

    p = sub.add_parser("finish-fast")
    add_bundle(p)
    p.add_argument("--format", default="md,docx,json")
    p.add_argument("--no-open", action="store_true")

    p = sub.add_parser("make-review-job")
    add_bundle(p)
    p.add_argument("--max-items", type=int, default=DEFAULT_QA_MAX_ITEMS)
    p.add_argument("--max-chars", type=int, default=DEFAULT_QA_MAX_CHARS)
    p.add_argument("--no-open", action="store_true")

    p = sub.add_parser("apply-qa")
    add_bundle(p)

    p = sub.add_parser("apply-review")
    add_bundle(p)

    p = sub.add_parser("render-final")
    add_bundle(p)
    p.add_argument("--format", default="md,docx,json")

    p = sub.add_parser("gate")
    add_bundle(p)
    p.add_argument("--json", action="store_true")

    p = sub.add_parser("status")
    add_bundle(p)
    p.add_argument("--json", action="store_true")

    p = sub.add_parser("make-repair-prompt")
    add_bundle(p)
    p.add_argument("--stage", choices=[TRANSLATION_STAGE, QA_STAGE], required=True)
    p.add_argument("--job-id", required=True)

    args = parser.parse_args()
    bundle = args.bundle.resolve()

    if args.command == "init-bundle":
        result = init_bundle(bundle, args.source_docx.resolve() if args.source_docx else None, args.source_input)
    elif args.command == "build-ir":
        quality_report = preflight_source(bundle, allow_layout_risk=args.allow_layout_risk)
        if quality_report["error_count"]:
            fail("source DOCX failed quality preflight; see logs/source_quality_report.json")
        result = build_ir(bundle, args.title)
        result = {"source_quality": quality_report, "build_ir": result}
    elif args.command == "make-translation-jobs":
        result = make_translation_jobs(bundle, args.max_items, args.max_chars, args.aux_max_items, args.aux_max_chars)
        if not args.no_open:
            open_stage_folders(bundle, TRANSLATION_STAGE)
    elif args.command == "prepare-translation":
        result = prepare_translation(
            bundle,
            args.source_docx.resolve(),
            args.source_input,
            args.title,
            args.max_items,
            args.max_chars,
            args.aux_max_items,
            args.aux_max_chars,
            not args.no_open,
            args.allow_layout_risk,
        )
    elif args.command == "preflight-source":
        result = preflight_source(
            bundle,
            args.source_docx.resolve() if args.source_docx else None,
            args.allow_layout_risk,
        )
    elif args.command == "open-hand-off":
        open_stage_folders(bundle, args.stage)
        result = {"opened": [str(stage_dir(bundle, args.stage)), str(stage_dir(bundle, args.stage) / "downloads")]}
    elif args.command == "accept-downloads":
        result = accept_downloads(bundle, args.stage, [path.resolve() for path in args.file])
    elif args.command == "validate-model-outputs":
        result = validate_model_outputs(bundle, args.stage)
    elif args.command == "merge-draft":
        result = merge_draft(bundle)
    elif args.command in {"make-qa-jobs", "make-review-job"}:
        result = make_qa_jobs(bundle, args.max_items, args.max_chars, not args.no_open)
    elif args.command == "finish-fast":
        result = finish_fast(bundle, args.format, not args.no_open)
    elif args.command in {"apply-qa", "apply-review"}:
        result = apply_qa(bundle)
    elif args.command == "render-final":
        result = render_final(bundle, args.format)
    elif args.command == "gate":
        result = gate(bundle)
    elif args.command == "status":
        result = status(bundle)
    elif args.command == "make-repair-prompt":
        result = make_repair_prompt(bundle, args.stage, args.job_id)
    else:  # pragma: no cover
        fail(f"unknown command {args.command}")

    print_result(result)
    if isinstance(result, dict) and result.get("error_count"):
        return 1
    if isinstance(result, dict) and result.get("gate_status") == "blocked":
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
