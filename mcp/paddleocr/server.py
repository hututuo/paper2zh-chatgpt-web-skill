#!/usr/bin/env python3
"""PaddleOCR MCP Server — OCR documents via PaddleOCR-VL-1.5 API."""

import base64
import json
import os
import re
import shutil
import tempfile
import time

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from mcp.server.fastmcp import FastMCP

DEFAULT_JOB_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
DEFAULT_MODEL = "PaddleOCR-VL-1.5"
TOKEN_ENV_NAMES = (
    "PADDLEOCR_ACCESS_TOKEN",
    "PADDLEOCR_MCP_AISTUDIO_ACCESS_TOKEN",
    "AISTUDIO_ACCESS_TOKEN",
)


def _load_env_file() -> None:
    """Load a small KEY=VALUE env file without adding a runtime dependency."""
    candidates = []
    explicit = os.environ.get("PADDLEOCR_MCP_ENV_FILE")
    if explicit:
        candidates.append(explicit)
    candidates.append(os.path.join(os.path.dirname(__file__), ".env"))

    for env_path in candidates:
        if not env_path or not os.path.exists(env_path):
            continue
        with open(env_path, "r", encoding="utf-8") as f:
            for raw_line in f:
                line = raw_line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value


_load_env_file()

JOB_URL = os.environ.get("PADDLEOCR_JOB_URL", DEFAULT_JOB_URL)
MODEL = os.environ.get("PADDLEOCR_MODEL", DEFAULT_MODEL)


def _access_token() -> str:
    for env_name in TOKEN_ENV_NAMES:
        token = os.environ.get(env_name, "").strip()
        if token:
            return token
    raise RuntimeError(
        "Missing PaddleOCR access token. Set PADDLEOCR_ACCESS_TOKEN or "
        "PADDLEOCR_MCP_AISTUDIO_ACCESS_TOKEN in the environment, or put it in "
        "a private .env file next to server.py."
    )


def _auth_headers() -> dict:
    return {"Authorization": f"bearer {_access_token()}"}

mcp = FastMCP("paddleocr")


# ---------------------------------------------------------------------------
# Core helpers
# ---------------------------------------------------------------------------

def _submit_job(file_path: str) -> dict:
    headers = _auth_headers()
    optional = {
        "useDocOrientationClassify": False,
        "useDocUnwarping": False,
        "useChartRecognition": False,
    }
    if file_path.startswith("http"):
        headers["Content-Type"] = "application/json"
        payload = {
            "fileUrl": file_path,
            "model": MODEL,
            "optionalPayload": optional,
        }
        resp = requests.post(JOB_URL, json=payload, headers=headers)
    else:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        data = {"model": MODEL, "optionalPayload": json.dumps(optional)}
        with open(file_path, "rb") as f:
            resp = requests.post(
                JOB_URL, headers=headers, data=data, files={"file": f}
            )
    if resp.status_code != 200:
        raise RuntimeError(f"API error {resp.status_code}: {resp.text}")
    return resp.json()


def _poll_job(job_id: str, timeout: int = 600) -> str:
    headers = _auth_headers()
    deadline = time.time() + timeout
    while time.time() < deadline:
        resp = requests.get(f"{JOB_URL}/{job_id}", headers=headers)
        if resp.status_code != 200:
            raise RuntimeError(f"Poll error {resp.status_code}: {resp.text}")
        data = resp.json()["data"]
        state = data["state"]
        if state == "done":
            return data["resultUrl"]["jsonUrl"]
        elif state == "failed":
            raise RuntimeError(f"Job failed: {data.get('errorMsg', 'unknown')}")
        time.sleep(5)
    raise TimeoutError(f"Job {job_id} did not finish within {timeout}s")


def _download_all_images(jsonl_url: str) -> dict:
    """Download all images from JSONL. Returns {img_path: bytes}."""
    resp = requests.get(jsonl_url)
    resp.raise_for_status()
    images = {}
    for line in resp.text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        result = json.loads(line)["result"]
        for res in result["layoutParsingResults"]:
            for img_path, img_url in res["markdown"].get("images", {}).items():
                if img_path not in images:
                    r = requests.get(img_url)
                    if r.status_code == 200:
                        images[img_path] = r.content
    return images


def _get_pages(jsonl_url: str):
    """Yield (md_text, images_dict) per page from JSONL."""
    resp = requests.get(jsonl_url)
    resp.raise_for_status()
    for line in resp.text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        result = json.loads(line)["result"]
        for res in result["layoutParsingResults"]:
            yield res["markdown"]["text"], res["markdown"].get("images", {})


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------

def _html_table_to_md(html_str: str) -> str:
    soup = BeautifulSoup(html_str, "html.parser")
    rows = soup.find_all("tr")
    if not rows:
        return html_str
    lines = []
    for ri, tr in enumerate(rows):
        cells = [c.get_text(strip=True).replace("\n", " ") for c in tr.find_all(["td", "th"])]
        lines.append("| " + " | ".join(cells) + " |")
        if ri == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(lines)


def _render_markdown(jsonl_url: str, output_dir: str) -> str:
    """Render OCR results to Markdown with images in a subfolder."""
    img_dir = os.path.join(output_dir, "images")
    os.makedirs(img_dir, exist_ok=True)

    # Download images
    all_images = _download_all_images(jsonl_url)
    for img_path, img_data in all_images.items():
        dest = os.path.join(img_dir, os.path.basename(img_path))
        with open(dest, "wb") as f:
            f.write(img_data)

    # Build markdown
    md_parts = []
    for md_text, _ in _get_pages(jsonl_url):
        # Fix image paths
        md_text = md_text.replace('src="imgs/', 'src="images/')
        # Convert HTML tables to markdown
        md_text = re.sub(
            r"<table[^>]*>.*?</table>",
            lambda m: _html_table_to_md(m.group(0)),
            md_text,
            flags=re.DOTALL,
        )
        md_parts.append(md_text)

    md_content = "\n\n---\n\n".join(md_parts)
    md_path = os.path.join(output_dir, os.path.basename(output_dir) + ".md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    return md_path


def _render_docx(jsonl_url: str, output_dir: str) -> str:
    """Render OCR results to Word with embedded images."""
    all_images = _download_all_images(jsonl_url)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for md_text, page_images in _get_pages(jsonl_url):
        # Download page-specific images for this page
        page_img_bytes = {}
        for img_path, img_url in page_images.items():
            if img_path in all_images:
                page_img_bytes[img_path] = all_images[img_path]

        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # HTML table
            if line.startswith("<table"):
                table_html = line
                while "</table>" not in table_html and i + 1 < len(lines):
                    i += 1
                    table_html += "\n" + lines[i]
                soup = BeautifulSoup(table_html, "html.parser")
                table_tag = soup.find("table")
                if table_tag:
                    rows = table_tag.find_all("tr")
                    max_cols = max(
                        (sum(int(c.get("colspan", 1)) for c in tr.find_all(["td", "th"])) for tr in rows),
                        default=0,
                    )
                    if max_cols > 0 and rows:
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = "Table Grid"
                        for ri, tr in enumerate(rows):
                            cols = tr.find_all(["td", "th"])
                            j = 0
                            for c in cols:
                                while j < max_cols and table.cell(ri, j).text:
                                    j += 1
                                if j >= max_cols:
                                    break
                                cell = table.cell(ri, j)
                                cell.text = c.get_text(strip=True)
                                colspan = int(c.get("colspan", 1))
                                if colspan > 1 and j + colspan <= max_cols:
                                    cell.merge(table.cell(ri, j + colspan - 1))
                                j += colspan
                i += 1
                continue

            # Image
            if "<img" in line and "src=" in line:
                img_match = re.search(r'src="([^"]+)"', line)
                if img_match:
                    img_key = img_match.group(1)
                    img_file = os.path.basename(img_key)
                    if img_key in page_img_bytes:
                        from io import BytesIO
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(BytesIO(page_img_bytes[img_key]), width=Cm(12))
                    else:
                        doc.add_paragraph(f"[Image: {img_file}]")
                i += 1
                continue

            # Skip raw html
            if line.startswith("<div") or line.startswith("</div"):
                i += 1
                continue

            # Headers
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=3)
            else:
                clean = re.sub(r"\$[^$]+\$", "", line).strip()
                if clean:
                    doc.add_paragraph(clean)
            i += 1

        doc.add_page_break()

    # Remove last empty page break
    docx_path = os.path.join(output_dir, os.path.basename(output_dir) + ".docx")
    doc.save(docx_path)
    return docx_path


# ---------------------------------------------------------------------------
# MCP tool
# ---------------------------------------------------------------------------

@mcp.tool()
def ocr_extract(
    file_path: str,
    output_format: str = "docx",
    output_dir: str = "",
) -> str:
    """Extract text from a document using PaddleOCR-VL-1.5.

    Args:
        file_path: Local file path or HTTP(S) URL to the document (PDF, image, etc.).
        output_format: Output format - "docx" (default, Word with embedded images),
                       "md" (Markdown with images folder), or "raw" (inline text).
        output_dir: Output directory. Defaults to a temp directory if empty.

    Returns:
        Path to the generated file(s), or raw markdown text if format is "raw".
    """
    resp = _submit_job(file_path)
    job_id = resp["data"]["jobId"]
    jsonl_url = _poll_job(job_id)

    # Determine output dir
    if not output_dir:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_dir = os.path.join(tempfile.gettempdir(), "paddleocr_" + base_name)
    os.makedirs(output_dir, exist_ok=True)

    if output_format == "md":
        md_path = _render_markdown(jsonl_url, output_dir)
        img_count = len(os.listdir(os.path.join(output_dir, "images")))
        return f"Done.\nMarkdown: {md_path}\nImages: {img_count} files in {output_dir}/images/"

    elif output_format == "docx":
        docx_path = _render_docx(jsonl_url, output_dir)
        return f"Done.\nWord: {docx_path}"

    else:  # raw
        pages = []
        for md_text, page_images in _get_pages(jsonl_url):
            # Inline base64
            for img_path, img_url in page_images.items():
                r = requests.get(img_url)
                if r.status_code == 200:
                    mime = r.headers.get("Content-Type", "image/jpeg")
                    b64 = base64.b64encode(r.content).decode()
                    md_text = md_text.replace(
                        f'src="{img_path}"', f'src="data:{mime};base64,{b64}"'
                    )
            pages.append(md_text)
        return f"Extracted {len(pages)} page(s).\n\n" + "\n\n---\n\n".join(pages)


if __name__ == "__main__":
    mcp.run()
